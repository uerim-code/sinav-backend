"""Sinav Otomasyon API — temiz isimlendirme."""

import io
from uuid import UUID
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from app.db.session import get_db
from app.models.sinav import (
    Fakulte, Program, Donem, Ders, Konu, Soru, SoruSecenegi,
    SoruSinavTipi, Cikabirlik, Sinav, SinavPlani, SinavSorusu,
    Ogrenci, Sonuc, OgrenciCevap, IliskiliSoruGrup, IliskiliSoru,
    DersKazanim, KazanimKonu, SoruKazanim, Bildirim,
)
from app.models.kullanici import Kullanici
from app.api.routes.deps import get_user, get_yetkili_user, get_admin_user, YETKILI_ROLLER
from app.core.mail import bildirim_maili

router = APIRouter(prefix="/sinav-oto", tags=["Sinav Otomasyon"])


# ── Bildirim helper ──
def bildirim_olustur(db: Session, kullanici_id: str, baslik: str, mesaj: str = "",
                     tip: str = "bilgi", link: str = None, mail_gonder: bool = True):
    """Uygulama ici bildirim olustur + opsiyonel e-posta gonder."""
    b = Bildirim(kullanici_id=kullanici_id, baslik=baslik, mesaj=mesaj, tip=tip, link=link)
    db.add(b)
    db.flush()
    if mail_gonder:
        user = db.query(Kullanici).filter_by(id=kullanici_id).first()
        if user and user.email:
            try:
                bildirim_maili(user.email, baslik, mesaj, link)
            except Exception:
                pass  # mail hata verse de bildirim kaydedilsin
    return b


def bildirim_tum_yetkililere(db: Session, baslik: str, mesaj: str = "",
                              tip: str = "bilgi", link: str = None, haric_id: str = None):
    """Tum admin ve ogretim uyelerine bildirim gonder."""
    kullanicilar = db.query(Kullanici).filter(
        Kullanici.aktif == True,
        Kullanici.rol.in_(list(YETKILI_ROLLER)),
    ).all()
    for u in kullanicilar:
        if haric_id and str(u.id) == haric_id:
            continue
        bildirim_olustur(db, str(u.id), baslik, mesaj, tip, link)



# ══════════════════════════════════════════════════════════════════
#  ISTATISTIK
# ══════════════════════════════════════════════════════════════════

@router.get("/istatistik")
def istatistik(db: Session = Depends(get_db), _=Depends(get_user)):
    return {
        "ders_sayisi": db.query(Ders).filter_by(aktif=True).count(),
        "konu_sayisi": db.query(Konu).count(),
        "soru_sayisi": db.query(Soru).count(),
        "sinav_sayisi": db.query(Sinav).count(),
    }


# ══════════════════════════════════════════════════════════════════
#  FAKULTE / PROGRAM / DONEM / DERS / KONU CRUD
# ══════════════════════════════════════════════════════════════════

# Fakulte
@router.get("/fakulteler")
def fakulte_listele(db: Session = Depends(get_db), _=Depends(get_user)):
    return [{"id": str(f.id), "ad": f.ad, "aktif": f.aktif}
            for f in db.query(Fakulte).filter_by(aktif=True).order_by(Fakulte.ad).all()]

@router.post("/fakulteler")
def fakulte_ekle(veri: dict, db: Session = Depends(get_db), _=Depends(get_admin_user)):
    f = Fakulte(ad=veri["ad"])
    db.add(f); db.commit(); db.refresh(f)
    return {"id": str(f.id), "ad": f.ad}

# Program
@router.get("/fakulteler/{fakulte_id}/programlar")
def program_listele(fakulte_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    return [{"id": str(p.id), "ad": p.ad, "aktif": p.aktif}
            for p in db.query(Program).filter_by(fakulte_id=str(fakulte_id), aktif=True).order_by(Program.ad).all()]

@router.post("/fakulteler/{fakulte_id}/programlar")
def program_ekle(fakulte_id: UUID, veri: dict, db: Session = Depends(get_db), _=Depends(get_admin_user)):
    p = Program(fakulte_id=str(fakulte_id), ad=veri["ad"])
    db.add(p); db.commit(); db.refresh(p)
    return {"id": str(p.id), "ad": p.ad}

# Donem
@router.get("/programlar/{program_id}/donemler")
def donem_listele(program_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    return [{"id": str(d.id), "ad": d.ad, "aktif": d.aktif}
            for d in db.query(Donem).filter_by(program_id=str(program_id), aktif=True).order_by(Donem.ad).all()]

@router.post("/programlar/{program_id}/donemler")
def donem_ekle(program_id: UUID, veri: dict, db: Session = Depends(get_db), _=Depends(get_admin_user)):
    d = Donem(program_id=str(program_id), ad=veri["ad"])
    db.add(d); db.commit(); db.refresh(d)
    return {"id": str(d.id), "ad": d.ad}

# Ders
@router.get("/donemler/{donem_id}/dersler")
def ders_listele(donem_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    return [{"id": str(d.id), "ad": d.ad, "haftalik_saat": d.haftalik_saat, "aktif": d.aktif}
            for d in db.query(Ders).filter_by(donem_id=str(donem_id), aktif=True).order_by(Ders.ad).all()]

@router.post("/donemler/{donem_id}/dersler")
def ders_ekle(donem_id: UUID, veri: dict, db: Session = Depends(get_db), _=Depends(get_admin_user)):
    d = Ders(donem_id=str(donem_id), ad=veri["ad"], haftalik_saat=veri.get("haftalik_saat", 0))
    db.add(d); db.commit(); db.refresh(d)
    return {"id": str(d.id), "ad": d.ad}

# Konu
@router.get("/dersler/{ders_id}/konular")
def konu_listele(ders_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    return [{"id": str(k.id), "ad": k.ad, "hafta": k.hafta, "saat": k.saat, "sira": k.sira,
             "soru_sayisi": db.query(Soru).filter_by(konu_id=str(k.id)).count()}
            for k in db.query(Konu).filter_by(ders_id=str(ders_id)).order_by(Konu.sira, Konu.hafta).all()]

@router.post("/dersler/{ders_id}/konular")
def konu_ekle(ders_id: UUID, veri: dict, db: Session = Depends(get_db), me=Depends(get_yetkili_user)):
    k = Konu(ders_id=str(ders_id), ad=veri["ad"], hafta=veri.get("hafta"), saat=veri.get("saat"),
             anlatan_id=str(me.id), sira=veri.get("sira", 0))
    db.add(k); db.commit(); db.refresh(k)
    return {"id": str(k.id), "ad": k.ad}

@router.delete("/konular/{konu_id}")
def konu_sil(konu_id: UUID, db: Session = Depends(get_db), _=Depends(get_yetkili_user)):
    db.query(Konu).filter_by(id=str(konu_id)).delete()
    db.commit()
    return {"ok": True}

@router.post("/dersler/{ders_id}/konular/toplu")
def konular_toplu_ekle(ders_id: UUID, veri: dict, db: Session = Depends(get_db), me=Depends(get_yetkili_user)):
    """Yapistirilan metinden konulari otomatik olustur.
    veri: { metin: str } veya { konular: [{ad, hafta, saat}] }
    Metin formati ornekleri:
      - 'Hafta 1: Konusu\nHafta 2: Konusu'
      - '1. Konusu\n2. Konusu'
      - 'Konusu (2 saat)\nKonusu (3 saat)'
    """
    import re
    konular_list = veri.get("konular")
    if not konular_list:
        metin = veri.get("metin", "").strip()
        if not metin:
            raise HTTPException(400, "Metin veya konular listesi gerekli")
        satirlar = [s.strip() for s in metin.split("\n") if s.strip()]
        konular_list = []
        for i, satir in enumerate(satirlar, 1):
            # "Hafta 1: Konu adi" veya "H1: Konu adi" veya "1. Hafta: Konu adi"
            m = re.match(r"^(?:Hafta|H|Week)\s*(\d+)\s*[:.\-–]\s*(.+)", satir, re.IGNORECASE)
            if m:
                konular_list.append({"hafta": int(m.group(1)), "ad": m.group(2).strip(), "saat": None})
                continue
            # "1. Konu adi" veya "1) Konu adi" veya "1- Konu adi"
            m = re.match(r"^(\d+)\s*[.):\-–]\s*(.+)", satir)
            if m:
                saat_m = re.search(r"\((\d+)\s*(?:saat|s|sa|hr|hour)\)", m.group(2), re.IGNORECASE)
                ad = re.sub(r"\(\d+\s*(?:saat|s|sa|hr|hour)\)", "", m.group(2)).strip()
                konular_list.append({"hafta": int(m.group(1)), "ad": ad, "saat": int(saat_m.group(1)) if saat_m else None})
                continue
            # Duz metin
            saat_m = re.search(r"\((\d+)\s*(?:saat|s|sa|hr|hour)\)", satir, re.IGNORECASE)
            ad = re.sub(r"\(\d+\s*(?:saat|s|sa|hr|hour)\)", "", satir).strip()
            if ad:
                konular_list.append({"hafta": i, "ad": ad, "saat": int(saat_m.group(1)) if saat_m else None})

    mevcut_sira = db.query(Konu).filter_by(ders_id=str(ders_id)).count()
    eklenen = 0
    for k in konular_list:
        mevcut_sira += 1
        db.add(Konu(ders_id=str(ders_id), ad=k["ad"], hafta=k.get("hafta"), saat=k.get("saat"),
                     anlatan_id=str(me.id), sira=mevcut_sira))
        eklenen += 1
    db.commit()
    return {"eklenen": eklenen}


# ══════════════════════════════════════════════════════════════════
#  SORU CRUD
# ══════════════════════════════════════════════════════════════════

@router.get("/konular/{konu_id}/sorular")
def soru_listele(konu_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    sorular = db.query(Soru).filter_by(konu_id=str(konu_id)).order_by(Soru.olusturuldu.desc()).all()
    sonuc = []
    for s in sorular:
        secenekler = db.query(SoruSecenegi).filter_by(soru_id=str(s.id)).order_by(SoruSecenegi.sira).all()
        kazanim_ids = [str(sk.kazanim_id) for sk in db.query(SoruKazanim).filter_by(soru_id=str(s.id)).all()]
        sonuc.append({
            "id": str(s.id), "soru_metni": s.soru_metni, "soru_tipi": s.soru_tipi,
            "bilgisel_duzey": s.bilgisel_duzey, "zorluk": s.zorluk,
            "kei": s.kei, "kgi": s.kgi, "cevaplama_suresi": s.cevaplama_suresi,
            "anahtar_kelimeler": s.anahtar_kelimeler or [],
            "kazanim_ids": kazanim_ids,
            "secenekler": [{"id": str(sc.id), "metin": sc.secenek_metni, "dogru": sc.dogru, "sira": sc.sira} for sc in secenekler],
        })
    return sonuc

@router.post("/sorular")
def soru_ekle(veri: dict, db: Session = Depends(get_db), me=Depends(get_yetkili_user)):
    s = Soru(
        konu_id=veri["konu_id"], soru_metni=veri["soru_metni"],
        soru_tipi=veri.get("soru_tipi", "tek_dogru"),
        bilgisel_duzey=veri.get("bilgisel_duzey", "Bilgi"),
        zorluk=veri.get("zorluk", "orta"),
        kei=veri.get("kei"), kgi=veri.get("kgi"),
        cevaplama_suresi=veri.get("cevaplama_suresi", 60),
        baslangic_tarihi=veri.get("baslangic_tarihi"),
        bitis_tarihi=veri.get("bitis_tarihi"),
        anahtar_kelimeler=veri.get("anahtar_kelimeler", []),
        kaynakca=veri.get("kaynakca"),
        yapilandirilmis_cevap=veri.get("yapilandirilmis_cevap"),
        olusturan_id=str(me.id),
    )
    db.add(s); db.flush()

    for sc in veri.get("secenekler", []):
        db.add(SoruSecenegi(soru_id=str(s.id), secenek_metni=sc["metin"], dogru=sc.get("dogru", False), sira=sc.get("sira", 0)))

    for tip in veri.get("sinav_tipleri", []):
        db.add(SoruSinavTipi(soru_id=str(s.id), sinav_tipi=tip))

    for tur in veri.get("cikabirlik", []):
        db.add(Cikabirlik(soru_id=str(s.id), tur=tur))

    for kid in veri.get("kazanim_ids", []):
        db.add(SoruKazanim(soru_id=str(s.id), kazanim_id=kid))

    db.commit()
    return {"id": str(s.id), "durum": "eklendi"}

@router.post("/sorular/toplu-yapistir")
def sorular_toplu_yapistir(veri: dict, db: Session = Depends(get_db), me=Depends(get_yetkili_user)):
    """Birden fazla soruyu tek seferde yapistirarak ekle.
    veri: { konu_id, sorular: [{soru_metni, secenekler: [{metin, dogru}], zorluk, bilgisel_duzey}] }
    """
    konu_id = veri.get("konu_id")
    if not konu_id:
        raise HTTPException(400, "konu_id gerekli")
    sorular_list = veri.get("sorular", [])
    if not sorular_list:
        raise HTTPException(400, "En az bir soru gerekli")

    eklenen = 0
    for soru_data in sorular_list:
        s = Soru(
            konu_id=konu_id, soru_metni=soru_data["soru_metni"],
            soru_tipi=soru_data.get("soru_tipi", "tek_dogru"),
            bilgisel_duzey=soru_data.get("bilgisel_duzey", "Bilgi"),
            zorluk=soru_data.get("zorluk", "orta"),
            cevaplama_suresi=soru_data.get("cevaplama_suresi", 60),
            olusturan_id=str(me.id),
        )
        db.add(s); db.flush()
        for i, sc in enumerate(soru_data.get("secenekler", [])):
            db.add(SoruSecenegi(soru_id=str(s.id), secenek_metni=sc["metin"], dogru=sc.get("dogru", False), sira=i))
        eklenen += 1
    db.commit()
    return {"eklenen": eklenen}

@router.put("/sorular/{soru_id}")
def soru_guncelle(soru_id: UUID, veri: dict, db: Session = Depends(get_db), _=Depends(get_yetkili_user)):
    """Mevcut soruyu guncelle — metin, secenekler, zorluk vb."""
    soru = db.query(Soru).filter_by(id=str(soru_id)).first()
    if not soru:
        raise HTTPException(404, "Soru bulunamadi")

    # Soru alanlarini guncelle
    if "soru_metni" in veri: soru.soru_metni = veri["soru_metni"]
    if "soru_tipi" in veri: soru.soru_tipi = veri["soru_tipi"]
    if "bilgisel_duzey" in veri: soru.bilgisel_duzey = veri["bilgisel_duzey"]
    if "zorluk" in veri: soru.zorluk = veri["zorluk"]
    if "cevaplama_suresi" in veri: soru.cevaplama_suresi = veri["cevaplama_suresi"]
    if "kei" in veri: soru.kei = veri["kei"]
    if "kgi" in veri: soru.kgi = veri["kgi"]
    if "anahtar_kelimeler" in veri: soru.anahtar_kelimeler = veri["anahtar_kelimeler"]
    if "kaynakca" in veri: soru.kaynakca = veri["kaynakca"]
    if "yapilandirilmis_cevap" in veri: soru.yapilandirilmis_cevap = veri["yapilandirilmis_cevap"]

    # Secenekleri guncelle
    if "secenekler" in veri:
        db.query(SoruSecenegi).filter_by(soru_id=str(soru_id)).delete()
        for sc in veri["secenekler"]:
            if sc.get("metin", "").strip():
                db.add(SoruSecenegi(
                    soru_id=str(soru_id), secenek_metni=sc["metin"],
                    dogru=sc.get("dogru", False), sira=sc.get("sira", 0),
                ))

    # Sinav tiplerini guncelle
    if "sinav_tipleri" in veri:
        db.query(SoruSinavTipi).filter_by(soru_id=str(soru_id)).delete()
        for tip in veri["sinav_tipleri"]:
            db.add(SoruSinavTipi(soru_id=str(soru_id), sinav_turu=tip))

    # Kazanimlari guncelle
    if "kazanim_ids" in veri:
        db.query(SoruKazanim).filter_by(soru_id=str(soru_id)).delete()
        for kid in veri["kazanim_ids"]:
            db.add(SoruKazanim(soru_id=str(soru_id), kazanim_id=kid))

    db.commit()
    return {"ok": True}


@router.delete("/sorular/{soru_id}")
def soru_sil(soru_id: UUID, db: Session = Depends(get_db), _=Depends(get_yetkili_user)):
    db.query(SoruSecenegi).filter_by(soru_id=str(soru_id)).delete()
    db.query(SoruSinavTipi).filter_by(soru_id=str(soru_id)).delete()
    db.query(Cikabirlik).filter_by(soru_id=str(soru_id)).delete()
    db.query(Soru).filter_by(id=str(soru_id)).delete()
    db.commit()
    return {"ok": True}


# ══════════════════════════════════════════════════════════════════
#  SINAV CRUD
# ══════════════════════════════════════════════════════════════════

@router.get("/sinavlar")
def sinav_listele(db: Session = Depends(get_db), _=Depends(get_user)):
    return [{"id": str(s.id), "ad": s.ad, "sinav_turu": s.sinav_turu, "durum": s.durum,
             "soru_sayisi": s.soru_sayisi, "baslangic": s.baslangic.isoformat() if s.baslangic else None}
            for s in db.query(Sinav).order_by(Sinav.baslangic.desc()).all()]

@router.post("/sinavlar")
def sinav_olustur(veri: dict, db: Session = Depends(get_db), me=Depends(get_yetkili_user)):
    s = Sinav(
        ders_id=veri.get("ders_id"), ad=veri["ad"],
        sinav_turu=veri.get("sinav_turu", "donem_ici"),
        sinav_kategorisi=veri.get("sinav_kategorisi", "kuramsal"),
        tam_puan=veri.get("tam_puan", 100), soru_sayisi=veri.get("soru_sayisi", 0),
        kitapcik_turu=veri.get("kitapcik_turu", "AB"),
        soru_secim_sekli=veri.get("soru_secim_sekli", "plan_ekraniyla"),
        olusturan_id=str(me.id), durum="taslak",
    )
    db.add(s); db.commit(); db.refresh(s)

    # Bildirim gonder
    bildirim_tum_yetkililere(
        db, baslik=f"Yeni sinav olusturuldu: {s.ad}",
        mesaj=f"{me.ad_soyad} tarafindan '{s.ad}' sinavi olusturuldu.",
        tip="basari", link="/sinavlar", haric_id=str(me.id),
    )
    db.commit()

    return {"id": str(s.id), "ad": s.ad}


@router.get("/sinavlar/{sinav_id}")
def sinav_detay(sinav_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    """Sinav detay bilgileri."""
    s = db.query(Sinav).filter_by(id=str(sinav_id)).first()
    if not s:
        raise HTTPException(404, "Sinav bulunamadi")
    ders = db.query(Ders).filter_by(id=str(s.ders_id)).first()
    soru_sayisi_atanmis = db.query(SinavSorusu).filter_by(sinav_id=str(sinav_id)).count()
    return {
        "id": str(s.id), "ad": s.ad, "sinav_turu": s.sinav_turu,
        "sinav_kategorisi": s.sinav_kategorisi, "durum": s.durum,
        "soru_sayisi": s.soru_sayisi, "tam_puan": s.tam_puan,
        "kitapcik_turu": s.kitapcik_turu,
        "soru_secim_sekli": s.soru_secim_sekli,
        "baslangic": s.baslangic.isoformat() if s.baslangic else None,
        "bitis": s.bitis.isoformat() if s.bitis else None,
        "ders_id": str(s.ders_id) if s.ders_id else None,
        "ders_ad": ders.ad if ders else None,
        "soru_sayisi_atanmis": soru_sayisi_atanmis,
    }


@router.put("/sinavlar/{sinav_id}")
def sinav_guncelle(sinav_id: UUID, veri: dict, db: Session = Depends(get_db), _=Depends(get_yetkili_user)):
    """Sinav bilgilerini guncelle."""
    s = db.query(Sinav).filter_by(id=str(sinav_id)).first()
    if not s:
        raise HTTPException(404, "Sinav bulunamadi")

    if "ad" in veri: s.ad = veri["ad"]
    if "sinav_turu" in veri: s.sinav_turu = veri["sinav_turu"]
    if "sinav_kategorisi" in veri: s.sinav_kategorisi = veri["sinav_kategorisi"]
    if "tam_puan" in veri: s.tam_puan = veri["tam_puan"]
    if "soru_sayisi" in veri: s.soru_sayisi = veri["soru_sayisi"]
    if "kitapcik_turu" in veri: s.kitapcik_turu = veri["kitapcik_turu"]
    if "durum" in veri: s.durum = veri["durum"]
    db.commit()
    return {"ok": True}


# ══════════════════════════════════════════════════════════════════
#  OPTIK OKUYUCU
# ══════════════════════════════════════════════════════════════════

@router.post("/optik-yukle/{sinav_id}")
async def optik_yukle(sinav_id: UUID, dosya: UploadFile = File(...), db: Session = Depends(get_db), _=Depends(get_yetkili_user)):
    """Optik okuyucu CSV/Excel yukle -> ogrenci cevaplari parse et -> kaydet.
    Format: ogrenci_no, ad_soyad, kitapcik, cevap1, cevap2, ...
    veya:   ogrenci_no, kitapcik, cevap1, cevap2, ...
    """
    import csv
    sinav = db.query(Sinav).filter_by(id=str(sinav_id)).first()
    if not sinav: raise HTTPException(404, "Sinav bulunamadi")

    icerik = await dosya.read()
    dosya_adi = dosya.filename or ""

    # Excel veya CSV parse
    rows = []
    if dosya_adi.endswith('.xlsx') or dosya_adi.endswith('.xls'):
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(icerik), read_only=True)
        ws = wb.active
        for row in ws.iter_rows(values_only=True):
            rows.append([str(c or '').strip() for c in row])
    else:
        try:
            text = icerik.decode('utf-8')
        except UnicodeDecodeError:
            text = icerik.decode('latin-1')
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)

    if len(rows) < 2:
        raise HTTPException(400, "Dosyada veri bulunamadi")

    # Baslik satirindan format algilama
    basliklar = [h.strip().lower() for h in rows[0]]
    has_ad = any('ad' in h or 'name' in h or 'soyad' in h for h in basliklar)
    # ad_soyad sutunu varsa: no, ad, kitapcik, cevaplar
    # yoksa: no, kitapcik, cevaplar
    cevap_baslangic = 3 if has_ad else 2

    sinav_sorulari = db.query(SinavSorusu).filter_by(sinav_id=str(sinav_id)).order_by(SinavSorusu.sira).all()
    soru_map = {s.sira: str(s.soru_id) for s in sinav_sorulari}

    # Onceki sonuclari temizle (yeniden yukleme)
    eski_sonuclar = db.query(Sonuc).filter_by(sinav_id=str(sinav_id)).all()
    for es in eski_sonuclar:
        db.query(OgrenciCevap).filter_by(sinav_sonucu_id=str(es.id)).delete()
    db.query(Sonuc).filter_by(sinav_id=str(sinav_id)).delete()
    db.query(Ogrenci).filter_by(sinav_id=str(sinav_id)).delete()
    db.flush()

    kaydedilen = 0
    for row in rows[1:]:
        if len(row) < cevap_baslangic + 1: continue
        ogrenci_no = row[0].strip()
        if not ogrenci_no: continue

        if has_ad:
            ad_soyad = row[1].strip()
            kitapcik = row[2].strip().upper() if len(row) > 2 else "A"
        else:
            ad_soyad = ""
            kitapcik = row[1].strip().upper() if len(row) > 1 else "A"

        if kitapcik not in "ABCDEFGH":
            kitapcik = "A"

        ogrenci = Ogrenci(sinav_id=str(sinav_id), ogrenci_no=ogrenci_no, ad=ad_soyad)
        db.add(ogrenci); db.flush()

        dogru_sayisi = 0; yanlis_sayisi = 0; bos_sayisi = 0
        for i, cevap in enumerate(row[cevap_baslangic:], 1):
            cevap = str(cevap).strip().upper()
            soru_id = soru_map.get(i)
            if not soru_id: continue

            dogru_sec = db.query(SoruSecenegi).filter_by(soru_id=soru_id, dogru=True).first()
            harfler = "ABCDE"
            dogru_harf = None
            if dogru_sec:
                secenekler = db.query(SoruSecenegi).filter_by(soru_id=soru_id).order_by(SoruSecenegi.sira).all()
                for idx, sc in enumerate(secenekler):
                    if sc.dogru and idx < len(harfler):
                        dogru_harf = harfler[idx]
                        break

            is_dogru = cevap == dogru_harf if dogru_harf and cevap else False
            if not cevap or cevap == '-' or cevap == '':
                bos_sayisi += 1
            elif is_dogru:
                dogru_sayisi += 1
            else:
                yanlis_sayisi += 1

            db.add(OgrenciCevap(
                sinav_sonucu_id="temp", soru_id=soru_id,
                verilen_secenek=cevap if cevap and cevap != '-' else None,
                dogru=is_dogru,
            ))

        toplam = dogru_sayisi + yanlis_sayisi + bos_sayisi
        net = dogru_sayisi - (yanlis_sayisi / 4) if toplam > 0 else 0
        ham_puan = (net / toplam * sinav.tam_puan) if toplam > 0 else 0

        sonuc = Sonuc(
            sinav_id=str(sinav_id), ogrenci_id=str(ogrenci.id),
            ham_puan=round(ham_puan, 2), net=round(net, 2),
            dogru=dogru_sayisi, yanlis=yanlis_sayisi, bos=bos_sayisi,
            yuzdelik=0, kitapcik=kitapcik,
        )
        db.add(sonuc); db.flush()

        # OgrenciCevap'lara dogru sinav_sonucu_id ata
        db.query(OgrenciCevap).filter_by(sinav_sonucu_id="temp").update({"sinav_sonucu_id": str(sonuc.id)})
        kaydedilen += 1

    db.commit()

    sonuclar = db.query(Sonuc).filter_by(sinav_id=str(sinav_id)).order_by(Sonuc.ham_puan.desc()).all()
    for idx, s in enumerate(sonuclar):
        s.yuzdelik = round((len(sonuclar) - idx) / len(sonuclar) * 100, 1)
    db.commit()

    # Bildirim gonder
    bildirim_tum_yetkililere(
        db, baslik=f"Optik sonuclar yuklendi: {sinav.ad}",
        mesaj=f"{kaydedilen} ogrenci sonucu yuklendi. Analiz sayfasindan detaylari inceleyebilirsiniz.",
        tip="basari", link="/analizler",
    )
    db.commit()

    # Otomatik zorluk guncelleme (madde analizi bazli)
    _zorluk_guncelle_from_analiz(db, str(sinav_id))

    return {"kaydedilen": kaydedilen, "toplam_satir": len(rows) - 1}


@router.get("/sinavlar/{sinav_id}/sonuc-sablon")
def sonuc_sablon_indir(sinav_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    """Sinava ozel ogrenci sonuc yukleme Excel sablonu."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    sinav = db.query(Sinav).filter_by(id=str(sinav_id)).first()
    if not sinav:
        raise HTTPException(404, "Sinav bulunamadi")

    sinav_sorulari = db.query(SinavSorusu).filter_by(sinav_id=str(sinav_id)).order_by(SinavSorusu.sira).all()
    soru_sayisi = len(sinav_sorulari) or sinav.soru_sayisi or 20

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Ogrenci Cevaplari"

    # Stiller
    baslik_font = Font(bold=True, color="FFFFFF", size=11)
    baslik_fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
    cevap_fill = PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid")
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    # Basliklar
    headers = ["Ogrenci No", "Ad Soyad", "Kitapcik"]
    for i in range(1, soru_sayisi + 1):
        headers.append(f"S{i}")

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = baslik_font
        cell.fill = baslik_fill
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_border

    # Sutun genislikleri
    ws.column_dimensions['A'].width = 14
    ws.column_dimensions['B'].width = 22
    ws.column_dimensions['C'].width = 10
    for i in range(4, 4 + soru_sayisi):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = 5

    # Ornek satirlar
    for r in range(2, 5):
        ws.cell(row=r, column=1, value=f"20210{r}001").border = thin_border
        ws.cell(row=r, column=2, value=f"Ornek Ogrenci {r - 1}").border = thin_border
        ws.cell(row=r, column=3, value="A").border = thin_border
        for c in range(4, 4 + soru_sayisi):
            cell = ws.cell(row=r, column=c, value="")
            cell.fill = cevap_fill
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border

    # Aciklama sayfasi
    ws2 = wb.create_sheet("Aciklama")
    aciklamalar = [
        ["OGRENCI SONUC YUKLEME SABLONU"],
        [""],
        [f"Sinav: {sinav.ad}"],
        [f"Soru Sayisi: {soru_sayisi}"],
        [""],
        ["SUTUN ACIKLAMALARI:"],
        ["Ogrenci No", "Ogrencinin numarasi (zorunlu)"],
        ["Ad Soyad", "Ogrencinin adi soyadi (opsiyonel ama onerilen)"],
        ["Kitapcik", "Sinav kitapcik harfi: A, B, C veya D"],
        ["S1, S2, ...", "Her soru icin ogrencinin verdigi cevap harfi: A, B, C, D veya E"],
        ["", "Bos birakilirsa veya - yazilirsa bos sayilir"],
        [""],
        ["ORNEK:"],
        ["20210001", "Ali Yilmaz", "A", "A", "B", "C", "D", "A", "B", "..."],
        ["20210002", "Ayse Demir", "B", "C", "A", "B", "E", "D", "A", "..."],
        [""],
        ["NOTLAR:"],
        ["- Ilk satir baslik satirdir, degistirmeyin"],
        ["- Ornek satirlari silip kendi verilerinizi girin"],
        ["- CSV formatinda da yukleyebilirsiniz"],
        ["- Yeniden yukleme yapilirsa onceki sonuclar silinir"],
    ]
    for r, row in enumerate(aciklamalar, 1):
        for c, val in enumerate(row, 1):
            cell = ws2.cell(row=r, column=c, value=val)
            if r == 1:
                cell.font = Font(bold=True, size=14)
            elif r == 6:
                cell.font = Font(bold=True, size=12)

    ws2.column_dimensions['A'].width = 18
    ws2.column_dimensions['B'].width = 50

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    from urllib.parse import quote
    filename = f"Sonuc_Sablonu_{sinav.ad.replace(' ', '_')}.xlsx"
    return StreamingResponse(
        buffer, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"}
    )


def _p_to_zorluk(p: float) -> str:
    """Zorluk indeksini (p) zorluk derecesine cevir."""
    if p < 0.20:
        return "cok_zor"
    elif p < 0.40:
        return "zor"
    elif p < 0.60:
        return "orta"
    elif p < 0.80:
        return "kolay"
    else:
        return "cok_kolay"


def _zorluk_guncelle_from_analiz(db: Session, sinav_id: str):
    """Sinav sonuclarindan zorluk indeksi hesaplayip sorularin zorluk derecesini guncelle."""
    sonuclar = db.query(Sonuc).filter_by(sinav_id=sinav_id).all()
    if not sonuclar:
        return

    sinav_sorulari = db.query(SinavSorusu).filter_by(sinav_id=sinav_id).all()
    n = len(sonuclar)
    if n == 0:
        return

    guncellenen = 0
    for ss in sinav_sorulari:
        dogru_sayisi = db.query(OgrenciCevap).filter_by(
            soru_id=str(ss.soru_id), dogru=True
        ).count()
        p = dogru_sayisi / n
        yeni_zorluk = _p_to_zorluk(p)

        soru = db.query(Soru).filter_by(id=str(ss.soru_id)).first()
        if soru and soru.zorluk != yeni_zorluk:
            soru.zorluk = yeni_zorluk
            guncellenen += 1

    if guncellenen > 0:
        db.commit()


# ══════════════════════════════════════════════════════════════════
#  ANALIZ ENDPOINTLERI
# ══════════════════════════════════════════════════════════════════

@router.get("/analiz/{sinav_id}/sinav")
def sinav_analizi(sinav_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    """Sinav genel istatistikleri + normallik testi."""
    import math
    sonuclar = db.query(Sonuc).filter_by(sinav_id=str(sinav_id)).all()
    if not sonuclar:
        raise HTTPException(404, "Henuz sonuc yok")

    puanlar = [s.ham_puan for s in sonuclar]
    n = len(puanlar)
    ortalama = sum(puanlar) / n
    sirali = sorted(puanlar)
    medyan = sirali[n // 2] if n % 2 == 1 else (sirali[n // 2 - 1] + sirali[n // 2]) / 2
    varyans = sum((p - ortalama) ** 2 for p in puanlar) / n if n > 0 else 0
    std_sapma = math.sqrt(varyans)
    carpiklik = sum((p - ortalama) ** 3 for p in puanlar) / (n * std_sapma ** 3) if std_sapma > 0 else 0
    basiklik = (sum((p - ortalama) ** 4 for p in puanlar) / (n * std_sapma ** 4)) - 3 if std_sapma > 0 else 0

    # Normallik testi (Shapiro-Wilk)
    shapiro_w, shapiro_p = None, None
    normallik_yorum = None
    try:
        from scipy import stats
        if 3 <= n <= 5000:
            w, p_val = stats.shapiro(puanlar)
            shapiro_w = round(w, 4)
            shapiro_p = round(p_val, 4)
            if p_val > 0.05:
                normallik_yorum = "Normal dagilim gosteriyor (p > 0.05)"
            else:
                normallik_yorum = "Normal dagilim gostermiyor (p ≤ 0.05)"
    except Exception:
        pass

    # Carpiklik/basiklik yorumu
    if abs(carpiklik) < 0.5 and abs(basiklik) < 1:
        dagilim_yorum = "Yaklasik simetrik dagilim"
    elif carpiklik > 0.5:
        dagilim_yorum = "Saga carpik — kolay sinav veya dusuk basari"
    elif carpiklik < -0.5:
        dagilim_yorum = "Sola carpik — zor sinav veya yuksek basari"
    else:
        dagilim_yorum = "Hafif asimetrik dagilim"

    sinav = db.query(Sinav).filter_by(id=str(sinav_id)).first()
    gecme = (sinav.tam_puan * 0.5) if sinav else 50
    basarili = sum(1 for p in puanlar if p >= gecme)
    sinirda = sum(1 for p in puanlar if gecme * 0.8 <= p < gecme)
    basarisiz = n - basarili - sinirda

    return {
        "ogrenci_sayisi": n, "ortalama": round(ortalama, 2), "medyan": round(medyan, 2),
        "std_sapma": round(std_sapma, 2), "carpiklik": round(carpiklik, 2),
        "basiklik": round(basiklik, 2),
        "min": round(min(puanlar), 2), "max": round(max(puanlar), 2),
        "basarili": basarili, "sinirda": sinirda, "basarisiz": basarisiz,
        "basari_orani": round(basarili / n * 100, 1) if n > 0 else 0,
        "basarili_yuzde": round(basarili / n * 100, 1) if n > 0 else 0,
        "sinirda_yuzde": round(sinirda / n * 100, 1) if n > 0 else 0,
        "basarisiz_yuzde": round(basarisiz / n * 100, 1) if n > 0 else 0,
        # Normallik testi
        "shapiro_w": shapiro_w,
        "shapiro_p": shapiro_p,
        "normallik_yorum": normallik_yorum,
        "dagilim_yorum": dagilim_yorum,
        # Yorumlar
        "yorumlar": {
            "ortalama": f"Sinav ortalamasi {round(ortalama, 1)} puan. " + (
                "Ortalama tam puanin %60 ustunde — genel basari iyi." if ortalama >= sinav.tam_puan * 0.6 else
                "Ortalama tam puanin %40-60 arasinda — orta duzey basari." if ortalama >= sinav.tam_puan * 0.4 else
                "Ortalama tam puanin %40 altinda — sinav cok zor veya ogrenci hazirlik yetersiz."
            ),
            "std_sapma": f"Standart sapma {round(std_sapma, 2)}. " + (
                "Dusuk yayilim — ogrenciler benzer performans gostermis, sinav ayirt edici degil." if std_sapma < sinav.tam_puan * 0.1 else
                "Orta yayilim — ogrenciler arasi makul farklilik var." if std_sapma < sinav.tam_puan * 0.2 else
                "Yuksek yayilim — ogrenciler arasi buyuk fark var, sinav iyi ayirt ediyor."
            ),
            "carpiklik": f"Carpiklik degeri {round(carpiklik, 2)}. " + (
                "Pozitif carpiklik — ogrencilerin cogu dusuk puan almis, sinav zor." if carpiklik > 0.5 else
                "Negatif carpiklik — ogrencilerin cogu yuksek puan almis, sinav kolay." if carpiklik < -0.5 else
                "Yaklasik simetrik dagilim — sinav zorlugu dengeli."
            ),
            "basiklik": f"Basiklik degeri {round(basiklik, 2)}. " + (
                "Sivri dagilim — puanlar ortalama etrafinda yogunlasmis." if basiklik > 1 else
                "Basik dagilim — puanlar genis bir araliga yayilmis." if basiklik < -1 else
                "Normal basiklik — tipik sinav dagilimi."
            ),
            "basari": f"Basarili: %{round(basarili / n * 100, 1) if n > 0 else 0}, Sinirda: %{round(sinirda / n * 100, 1) if n > 0 else 0}, Basarisiz: %{round(basarisiz / n * 100, 1) if n > 0 else 0}. " + (
                "Basari orani yuksek — sinav amacina uygun." if basarili / n > 0.6 else
                "Basari orani dusuk — sinav zorlugu veya ogretim yontemi gozden gecirilmeli." if basarili / n < 0.3 else
                "Basari orani orta duzey."
            ),
        },
    }


@router.get("/analiz/{sinav_id}/madde")
def madde_analizi(sinav_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    """Madde (soru) analizi — zorluk ve ayiricilik indeksi."""
    sonuclar = db.query(Sonuc).filter_by(sinav_id=str(sinav_id)).order_by(Sonuc.ham_puan.desc()).all()
    if not sonuclar:
        raise HTTPException(404, "Henuz sonuc yok")

    sinav_sorulari = db.query(SinavSorusu).filter_by(sinav_id=str(sinav_id)).order_by(SinavSorusu.sira).all()
    n = len(sonuclar)
    ust_grup = set(str(s.ogrenci_id) for s in sonuclar[:n // 3])
    alt_grup = set(str(s.ogrenci_id) for s in sonuclar[-(n // 3):])

    maddeler = []
    for ss in sinav_sorulari:
        dogru_cevaplar = db.query(OgrenciCevap).filter_by(soru_id=str(ss.soru_id), dogru=True).all()
        dogru_ids = {str(c.sinav_sonucu_id) for c in dogru_cevaplar}

        sonuc_ogrenci = {str(s.id): str(s.ogrenci_id) for s in sonuclar}

        toplam_dogru = len(dogru_ids)
        ust_dogru = sum(1 for sid in dogru_ids if sonuc_ogrenci.get(sid) in ust_grup)
        alt_dogru = sum(1 for sid in dogru_ids if sonuc_ogrenci.get(sid) in alt_grup)

        p = toplam_dogru / n if n > 0 else 0
        ust_n = len(ust_grup) or 1
        D = (ust_dogru / ust_n) - (alt_dogru / (len(alt_grup) or 1))

        durum = "iyi" if p >= 0.3 and p <= 0.7 and D >= 0.3 else ("sorunlu" if D < 0.2 else "dikkat")

        maddeler.append({
            "soru_no": ss.sira, "soru_id": str(ss.soru_id),
            "zorluk": round(p, 3), "ayiricilik": round(D, 3), "durum": durum,
        })

    return maddeler


@router.get("/analiz/{sinav_id}/madde-bilgi")
def madde_bilgi_fonksiyonu(sinav_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    """Madde Bilgi Fonksiyonu (IIF) — 2PL IRT yaklasimi.
    Her maddenin hangi yetenek duzeyinde en cok bilgi verdigini hesaplar."""
    import math

    sonuclar = db.query(Sonuc).filter_by(sinav_id=str(sinav_id)).order_by(Sonuc.ham_puan.desc()).all()
    if not sonuclar:
        raise HTTPException(404, "Henuz sonuc yok")

    sinav_sorulari = db.query(SinavSorusu).filter_by(sinav_id=str(sinav_id)).order_by(SinavSorusu.sira).all()
    n = len(sonuclar)
    ust_sinir = max(n // 3, 1)
    ust_grup = set(str(s.id) for s in sonuclar[:ust_sinir])
    alt_grup = set(str(s.id) for s in sonuclar[-ust_sinir:])

    # Theta araligi: -3 ile +3 arasi (standart normal)
    theta_range = [round(-3 + i * 0.25, 2) for i in range(25)]  # -3, -2.75, ..., +3

    maddeler = []
    test_bilgi = [0.0] * len(theta_range)  # Test bilgi fonksiyonu (toplam)

    for ss in sinav_sorulari:
        dogru_cevaplar = db.query(OgrenciCevap).filter_by(soru_id=str(ss.soru_id), dogru=True).all()
        dogru_ids = {str(c.sinav_sonucu_id) for c in dogru_cevaplar}

        sonuc_ogrenci = {str(s.id): str(s.ogrenci_id) for s in sonuclar}
        toplam_dogru = len(dogru_ids)
        ust_dogru = sum(1 for sid in dogru_ids if sid in ust_grup)
        alt_dogru = sum(1 for sid in dogru_ids if sid in alt_grup)

        p = toplam_dogru / n if n > 0 else 0.5
        D = (ust_dogru / ust_sinir) - (alt_dogru / ust_sinir) if ust_sinir > 0 else 0

        # IRT 2PL parametreleri yaklasik tahmin
        # b (zorluk) = logit donusumu: b = -ln(p / (1-p)) / 1.7
        # a (ayiricilik) = D'den yaklasik: a = D * 2.5
        p_clamp = max(0.01, min(0.99, p))
        b = -math.log(p_clamp / (1 - p_clamp)) / 1.7
        a = max(0.1, D * 2.5)

        # Madde bilgi fonksiyonu: I(θ) = a² × P(θ) × Q(θ)
        # P(θ) = 1 / (1 + exp(-1.7 × a × (θ - b)))
        bilgi_degerleri = []
        max_bilgi = 0
        max_theta = 0
        for ti, theta in enumerate(theta_range):
            z = 1.7 * a * (theta - b)
            z_clamp = max(-30, min(30, z))
            P = 1 / (1 + math.exp(-z_clamp))
            Q = 1 - P
            info = a * a * P * Q
            bilgi_degerleri.append(round(info, 4))
            test_bilgi[ti] += info
            if info > max_bilgi:
                max_bilgi = info
                max_theta = theta

        maddeler.append({
            "soru_no": ss.sira,
            "soru_id": str(ss.soru_id),
            "a": round(a, 3),  # ayiricilik parametresi
            "b": round(b, 3),  # zorluk parametresi
            "p": round(p, 3),
            "D": round(D, 3),
            "max_bilgi": round(max_bilgi, 4),
            "max_theta": max_theta,
            "bilgi": bilgi_degerleri,
        })

    return {
        "theta_range": theta_range,
        "maddeler": maddeler,
        "test_bilgi": [round(v, 4) for v in test_bilgi],
        "ogrenci_sayisi": n,
        "madde_sayisi": len(sinav_sorulari),
    }


@router.get("/analiz/{sinav_id}/celdirici")
def celdirici_analizi(sinav_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    """Celdirici analizi — her secenek icin secilme orani ve ust/alt grup dagilimi."""
    sonuclar = db.query(Sonuc).filter_by(sinav_id=str(sinav_id)).order_by(Sonuc.ham_puan.desc()).all()
    if not sonuclar:
        raise HTTPException(404, "Henuz sonuc yok")

    sinav_sorulari = db.query(SinavSorusu).filter_by(sinav_id=str(sinav_id)).order_by(SinavSorusu.sira).all()
    n = len(sonuclar)
    ust_sinir = max(n // 3, 1)
    ust_sonuc_ids = {str(s.id) for s in sonuclar[:ust_sinir]}
    alt_sonuc_ids = {str(s.id) for s in sonuclar[-ust_sinir:]}
    harfler = list("ABCDE")

    sonuc_liste = []
    for ss in sinav_sorulari:
        soru = db.query(Soru).filter_by(id=str(ss.soru_id)).first()
        secenekler_db = db.query(SoruSecenegi).filter_by(soru_id=str(ss.soru_id)).order_by(SoruSecenegi.sira).all()
        cevaplar = db.query(OgrenciCevap).filter_by(soru_id=str(ss.soru_id)).all()

        # Dogru cevap harfini bul
        dogru_harf = None
        for idx, sc in enumerate(secenekler_db):
            if sc.dogru and idx < len(harfler):
                dogru_harf = harfler[idx]
                break

        # Her secenek icin sayimlari hesapla
        secenek_sayim = {h: {"toplam": 0, "ust": 0, "alt": 0} for h in harfler[:len(secenekler_db)]}
        secenek_sayim["BOS"] = {"toplam": 0, "ust": 0, "alt": 0}

        for c in cevaplar:
            harf = (c.verilen_secenek or "").strip().upper()
            if not harf or harf == "-":
                key = "BOS"
            elif harf in secenek_sayim:
                key = harf
            else:
                key = "BOS"

            secenek_sayim[key]["toplam"] += 1
            if str(c.sinav_sonucu_id) in ust_sonuc_ids:
                secenek_sayim[key]["ust"] += 1
            if str(c.sinav_sonucu_id) in alt_sonuc_ids:
                secenek_sayim[key]["alt"] += 1

        # Secenek detaylari olustur
        secenekler_detay = []
        for idx, sc in enumerate(secenekler_db):
            if idx >= len(harfler):
                break
            h = harfler[idx]
            s_data = secenek_sayim.get(h, {"toplam": 0, "ust": 0, "alt": 0})
            oran = round(s_data["toplam"] / n * 100, 1) if n > 0 else 0
            ust_oran = round(s_data["ust"] / ust_sinir * 100, 1) if ust_sinir > 0 else 0
            alt_oran = round(s_data["alt"] / ust_sinir * 100, 1) if ust_sinir > 0 else 0

            # Celdirici etkinligi degerlendirmesi
            dogru_mu = h == dogru_harf
            if dogru_mu:
                etkinlik = "dogru"
            elif s_data["toplam"] == 0:
                etkinlik = "etkisiz"  # hic secilmemis
            elif alt_oran > ust_oran:
                etkinlik = "etkili"  # alt gruptan cok seciliyorsa iyi celdirici
            elif oran < 5:
                etkinlik = "etkisiz"  # %5'ten az secilmis
            else:
                etkinlik = "zayif"  # ust gruptan da cok seciliyorsa kotu

            secenekler_detay.append({
                "harf": h,
                "metin": sc.secenek_metni,
                "dogru": dogru_mu,
                "toplam": s_data["toplam"],
                "oran": oran,
                "ust_grup": s_data["ust"],
                "ust_oran": ust_oran,
                "alt_grup": s_data["alt"],
                "alt_oran": alt_oran,
                "etkinlik": etkinlik,
            })

        # Bos cevap
        bos = secenek_sayim["BOS"]
        bos_oran = round(bos["toplam"] / n * 100, 1) if n > 0 else 0

        sonuc_liste.append({
            "soru_no": ss.sira,
            "soru_id": str(ss.soru_id),
            "soru_metni": soru.soru_metni[:100] if soru else "",
            "dogru_harf": dogru_harf,
            "secenekler": secenekler_detay,
            "bos_sayisi": bos["toplam"],
            "bos_oran": bos_oran,
        })

    return {
        "ogrenci_sayisi": n,
        "ust_grup_n": ust_sinir,
        "alt_grup_n": ust_sinir,
        "sorular": sonuc_liste,
    }


@router.post("/analiz/{sinav_id}/zorluk-guncelle")
def zorluk_otomatik_guncelle(sinav_id: UUID, db: Session = Depends(get_db), _=Depends(get_yetkili_user)):
    """Madde analizi sonuclarina gore soru zorluk derecelerini otomatik guncelle."""
    sonuclar = db.query(Sonuc).filter_by(sinav_id=str(sinav_id)).all()
    if not sonuclar:
        raise HTTPException(404, "Henuz sonuc yok")

    sinav_sorulari = db.query(SinavSorusu).filter_by(sinav_id=str(sinav_id)).all()
    n = len(sonuclar)
    guncellenen = 0
    detay = []

    for ss in sinav_sorulari:
        dogru_sayisi = db.query(OgrenciCevap).filter_by(soru_id=str(ss.soru_id), dogru=True).count()
        p = dogru_sayisi / n
        yeni_zorluk = _p_to_zorluk(p)

        soru = db.query(Soru).filter_by(id=str(ss.soru_id)).first()
        if soru:
            eski = soru.zorluk
            if eski != yeni_zorluk:
                soru.zorluk = yeni_zorluk
                guncellenen += 1
            detay.append({
                "soru_no": ss.sira, "p": round(p, 3),
                "eski_zorluk": eski, "yeni_zorluk": yeni_zorluk,
                "degisti": eski != yeni_zorluk,
            })

    db.commit()
    return {"guncellenen": guncellenen, "toplam": len(sinav_sorulari), "detay": detay}


@router.get("/analiz/{sinav_id}/guvenirlik")
def guvenirlik_analizi(sinav_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    """Sinav guvenirlik analizi: Cronbach Alpha, KR-20, KR-21, SEM."""
    import math

    sonuclar = db.query(Sonuc).filter_by(sinav_id=str(sinav_id)).order_by(Sonuc.ham_puan.desc()).all()
    if not sonuclar:
        raise HTTPException(404, "Henuz sonuc yok")

    sinav = db.query(Sinav).filter_by(id=str(sinav_id)).first()
    sinav_sorulari = db.query(SinavSorusu).filter_by(sinav_id=str(sinav_id)).order_by(SinavSorusu.sira).all()
    n_ogrenci = len(sonuclar)
    k = len(sinav_sorulari)  # madde sayisi

    if k < 2 or n_ogrenci < 2:
        raise HTTPException(400, "Guvenirlik hesabi icin en az 2 soru ve 2 ogrenci gerekli")

    # Ogrenci-soru matrisi olustur (1=dogru, 0=yanlis)
    sonuc_map = {str(s.id): s for s in sonuclar}
    soru_ids = [str(ss.soru_id) for ss in sinav_sorulari]
    matris = []  # [ogrenci][soru] = 0 veya 1

    for sonuc in sonuclar:
        cevaplar = db.query(OgrenciCevap).filter_by(sinav_sonucu_id=str(sonuc.id)).all()
        cevap_map = {str(c.soru_id): (1 if c.dogru else 0) for c in cevaplar}
        satir = [cevap_map.get(sid, 0) for sid in soru_ids]
        matris.append(satir)

    # --- Cronbach Alpha ---
    # alpha = (k / (k-1)) * (1 - sum(si^2) / st^2)
    # si^2 = her maddenin varyansı, st^2 = toplam puanın varyansı
    toplam_puanlar = [sum(satir) for satir in matris]
    ortalama_toplam = sum(toplam_puanlar) / n_ogrenci
    varyans_toplam = sum((x - ortalama_toplam) ** 2 for x in toplam_puanlar) / n_ogrenci

    madde_varyanslari = []
    madde_p_degerleri = []
    for j in range(k):
        madde_puanlari = [matris[i][j] for i in range(n_ogrenci)]
        ort = sum(madde_puanlari) / n_ogrenci
        var_j = sum((x - ort) ** 2 for x in madde_puanlari) / n_ogrenci
        madde_varyanslari.append(var_j)
        madde_p_degerleri.append(ort)  # p = dogru orani

    sum_madde_varyans = sum(madde_varyanslari)
    cronbach_alpha = (k / (k - 1)) * (1 - sum_madde_varyans / varyans_toplam) if varyans_toplam > 0 else 0

    # --- KR-20 ---
    # KR-20 = (k / (k-1)) * (1 - sum(p*q) / st^2)
    # Dikotomik veriler icin Cronbach Alpha ile esit
    sum_pq = sum(p * (1 - p) for p in madde_p_degerleri)
    kr20 = (k / (k - 1)) * (1 - sum_pq / varyans_toplam) if varyans_toplam > 0 else 0

    # --- KR-21 ---
    # KR-21 = (k / (k-1)) * (1 - (ortalama * (k - ortalama)) / (k * st^2))
    kr21 = (k / (k - 1)) * (1 - (ortalama_toplam * (k - ortalama_toplam)) / (k * varyans_toplam)) if varyans_toplam > 0 else 0

    # --- SEM (Standart Hata) ---
    std_sapma = math.sqrt(varyans_toplam) if varyans_toplam > 0 else 0
    sem = std_sapma * math.sqrt(1 - cronbach_alpha) if cronbach_alpha < 1 else 0

    # --- Nokta Biserial Korelasyon (her madde icin) ---
    madde_rpb = []
    for j in range(k):
        dogru_puanlar = [toplam_puanlar[i] for i in range(n_ogrenci) if matris[i][j] == 1]
        yanlis_puanlar = [toplam_puanlar[i] for i in range(n_ogrenci) if matris[i][j] == 0]

        if len(dogru_puanlar) == 0 or len(yanlis_puanlar) == 0:
            madde_rpb.append({"soru_no": sinav_sorulari[j].sira, "rpb": 0})
            continue

        ort_dogru = sum(dogru_puanlar) / len(dogru_puanlar)
        ort_yanlis = sum(yanlis_puanlar) / len(yanlis_puanlar)
        p_j = madde_p_degerleri[j]
        q_j = 1 - p_j

        rpb = ((ort_dogru - ort_yanlis) / std_sapma) * math.sqrt(p_j * q_j) if std_sapma > 0 and p_j > 0 and q_j > 0 else 0
        madde_rpb.append({"soru_no": sinav_sorulari[j].sira, "soru_id": str(sinav_sorulari[j].soru_id), "rpb": round(rpb, 3)})

    # Yorum
    if cronbach_alpha >= 0.90:
        yorum = "Cok yuksek guvenirlik — mukemmel"
    elif cronbach_alpha >= 0.80:
        yorum = "Yuksek guvenirlik — iyi"
    elif cronbach_alpha >= 0.70:
        yorum = "Kabul edilebilir guvenirlik"
    elif cronbach_alpha >= 0.60:
        yorum = "Sinirda guvenirlik — iyilestirme gerekli"
    else:
        yorum = "Dusuk guvenirlik — sinav gozden gecirilmeli"

    return {
        "cronbach_alpha": round(cronbach_alpha, 4),
        "kr20": round(kr20, 4),
        "kr21": round(kr21, 4),
        "sem": round(sem, 3),
        "std_sapma": round(std_sapma, 3),
        "ortalama": round(ortalama_toplam, 2),
        "madde_sayisi": k,
        "ogrenci_sayisi": n_ogrenci,
        "yorum": yorum,
        "madde_rpb": madde_rpb,
    }


@router.get("/analiz/{sinav_id}/ogrenci")
def ogrenci_analizi(sinav_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    """Ogrenci bazli sonuclar + Z-puani + T-puani."""
    import math
    sonuclar = db.query(Sonuc).filter_by(sinav_id=str(sinav_id)).order_by(Sonuc.ham_puan.desc()).all()
    if not sonuclar:
        raise HTTPException(404, "Henuz sonuc yok")

    puanlar = [s.ham_puan for s in sonuclar]
    n = len(puanlar)
    ortalama = sum(puanlar) / n
    varyans = sum((p - ortalama) ** 2 for p in puanlar) / n
    std_sapma = math.sqrt(varyans) if varyans > 0 else 1

    liste = []
    for idx, s in enumerate(sonuclar, 1):
        ogrenci = db.query(Ogrenci).filter_by(id=str(s.ogrenci_id)).first()
        # Z-puani = (X - ortalama) / std_sapma
        z_puan = round((s.ham_puan - ortalama) / std_sapma, 2) if std_sapma > 0 else 0
        # T-puani = 50 + 10 * Z
        t_puan = round(50 + 10 * z_puan, 1)
        # Basari yorumu
        if z_puan >= 1.5:
            basari_yorum = "Cok basarili"
        elif z_puan >= 0.5:
            basari_yorum = "Basarili"
        elif z_puan >= -0.5:
            basari_yorum = "Ortalama"
        elif z_puan >= -1.5:
            basari_yorum = "Basari duzeyi dusuk"
        else:
            basari_yorum = "Cok dusuk basari"

        liste.append({
            "sira": idx, "ogrenci_no": ogrenci.ogrenci_no if ogrenci else "?",
            "ad": ogrenci.ad if ogrenci else "?",
            "puan": s.ham_puan, "net": s.net, "dogru": s.dogru,
            "yanlis": s.yanlis, "bos": s.bos, "yuzdelik": s.yuzdelik,
            "kitapcik": s.kitapcik,
            "z_puan": z_puan, "t_puan": t_puan, "basari_yorum": basari_yorum,
        })
    return liste


@router.get("/analiz/{sinav_id}/ogrenci-profil")
def ogrenci_basari_profili(sinav_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    """Ogrenci basari profili — konu bazli guclu/zayif alan analizi."""
    sonuclar = db.query(Sonuc).filter_by(sinav_id=str(sinav_id)).order_by(Sonuc.ham_puan.desc()).all()
    if not sonuclar:
        raise HTTPException(404, "Henuz sonuc yok")

    sinav_sorulari = db.query(SinavSorusu).filter_by(sinav_id=str(sinav_id)).order_by(SinavSorusu.sira).all()

    # Soru -> konu eslestirmesi
    soru_konu_map = {}  # soru_id -> {konu_id, konu_ad}
    konu_set = {}  # konu_id -> konu_ad
    for ss in sinav_sorulari:
        soru = db.query(Soru).filter_by(id=str(ss.soru_id)).first()
        if soru:
            konu = db.query(Konu).filter_by(id=str(soru.konu_id)).first()
            if konu:
                soru_konu_map[str(ss.soru_id)] = {"konu_id": str(konu.id), "konu_ad": konu.ad, "hafta": konu.hafta}
                konu_set[str(konu.id)] = {"ad": konu.ad, "hafta": konu.hafta}

    konular = sorted(konu_set.items(), key=lambda x: x[1].get("hafta") or 0)

    # Konu bazli sinav geneli istatistikleri
    konu_genel = {}  # konu_id -> {toplam_soru, toplam_dogru}
    for kid, kinfo in konular:
        konu_genel[kid] = {"toplam_soru": 0, "toplam_dogru": 0}

    # Her ogrenci icin konu bazli basari
    ogrenci_profiller = []
    for sonuc in sonuclar:
        ogrenci = db.query(Ogrenci).filter_by(id=str(sonuc.ogrenci_id)).first()
        cevaplar = db.query(OgrenciCevap).filter_by(sinav_sonucu_id=str(sonuc.id)).all()

        konu_basari = {}  # konu_id -> {dogru, yanlis, toplam}
        for kid, _ in konular:
            konu_basari[kid] = {"dogru": 0, "yanlis": 0, "toplam": 0}

        for c in cevaplar:
            skm = soru_konu_map.get(str(c.soru_id))
            if not skm:
                continue
            kid = skm["konu_id"]
            if kid not in konu_basari:
                konu_basari[kid] = {"dogru": 0, "yanlis": 0, "toplam": 0}
            konu_basari[kid]["toplam"] += 1
            if c.dogru:
                konu_basari[kid]["dogru"] += 1
                konu_genel[kid]["toplam_dogru"] += 1
            else:
                konu_basari[kid]["yanlis"] += 1
            konu_genel[kid]["toplam_soru"] += 1

        # Guclu ve zayif konulari belirle
        konu_detay = []
        for kid, kinfo in konular:
            kb = konu_basari.get(kid, {"dogru": 0, "yanlis": 0, "toplam": 0})
            oran = round(kb["dogru"] / kb["toplam"] * 100, 1) if kb["toplam"] > 0 else 0
            durum = "guclu" if oran >= 70 else ("zayif" if oran < 40 else "orta")
            konu_detay.append({
                "konu_id": kid, "konu_ad": kinfo["ad"], "hafta": kinfo.get("hafta"),
                "dogru": kb["dogru"], "yanlis": kb["yanlis"], "toplam": kb["toplam"],
                "oran": oran, "durum": durum,
            })

        guclu = [k["konu_ad"] for k in konu_detay if k["durum"] == "guclu"]
        zayif = [k["konu_ad"] for k in konu_detay if k["durum"] == "zayif"]

        ogrenci_profiller.append({
            "ogrenci_no": ogrenci.ogrenci_no if ogrenci else "?",
            "ad": ogrenci.ad if ogrenci else "?",
            "puan": sonuc.ham_puan,
            "konular": konu_detay,
            "guclu_konular": guclu,
            "zayif_konular": zayif,
        })

    # Konu bazli genel basari ozeti
    n = len(sonuclar)
    konu_ozet = []
    for kid, kinfo in konular:
        kg = konu_genel.get(kid, {"toplam_soru": 0, "toplam_dogru": 0})
        oran = round(kg["toplam_dogru"] / kg["toplam_soru"] * 100, 1) if kg["toplam_soru"] > 0 else 0
        konu_ozet.append({
            "konu_id": kid, "konu_ad": kinfo["ad"], "hafta": kinfo.get("hafta"),
            "basari_orani": oran,
            "yorum": "Basarili" if oran >= 70 else ("Iyilestirme gerekli" if oran < 40 else "Orta"),
        })

    return {
        "ogrenci_sayisi": n,
        "konu_ozet": konu_ozet,
        "ogrenciler": ogrenci_profiller,
    }


@router.get("/analiz/{sinav_id}/kazanim")
def kazanim_analizi(sinav_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    """Ogrenim Kazanimi analizi — basari orani, uyum matrisi, erisim duzeyi, Bloom dagilimi."""
    sonuclar = db.query(Sonuc).filter_by(sinav_id=str(sinav_id)).order_by(Sonuc.ham_puan.desc()).all()
    sinav_sorulari = db.query(SinavSorusu).filter_by(sinav_id=str(sinav_id)).order_by(SinavSorusu.sira).all()
    sinav = db.query(Sinav).filter_by(id=str(sinav_id)).first()
    if not sinav:
        raise HTTPException(404, "Sinav bulunamadi")

    n = len(sonuclar)

    # Ders kazanimlarini al
    ders_kazanimlar = db.query(DersKazanim).filter_by(ders_id=str(sinav.ders_id)).order_by(DersKazanim.sira).all()
    if not ders_kazanimlar:
        return {"kazanimlar": [], "uyum_matrisi": [], "bloom_dagilimi": {}, "yorum": "Bu derse ait kazanim tanimlanmamis."}

    # Soru -> kazanim eslestirmesi
    soru_kazanim_map = {}  # soru_id -> [kazanim_id, ...]
    for ss in sinav_sorulari:
        sk_list = db.query(SoruKazanim).filter_by(soru_id=str(ss.soru_id)).all()
        soru_kazanim_map[str(ss.soru_id)] = [str(sk.kazanim_id) for sk in sk_list]

    # Soru -> bilgisel duzey
    soru_bloom_map = {}
    for ss in sinav_sorulari:
        soru = db.query(Soru).filter_by(id=str(ss.soru_id)).first()
        if soru:
            soru_bloom_map[str(ss.soru_id)] = soru.bilgisel_duzey or "bilgi"

    # 16. Kazanim basari orani
    kazanim_sonuc = {}
    for kaz in ders_kazanimlar:
        kazanim_sonuc[str(kaz.id)] = {
            "kod": kaz.kod, "aciklama": kaz.aciklama, "bloom": kaz.bloom_duzeyi,
            "soru_sayisi": 0, "toplam_dogru": 0, "toplam_cevap": 0,
        }

    for ss in sinav_sorulari:
        kaz_ids = soru_kazanim_map.get(str(ss.soru_id), [])
        if not kaz_ids:
            continue
        dogru_count = db.query(OgrenciCevap).filter_by(soru_id=str(ss.soru_id), dogru=True).count() if n > 0 else 0
        for kid in kaz_ids:
            if kid in kazanim_sonuc:
                kazanim_sonuc[kid]["soru_sayisi"] += 1
                kazanim_sonuc[kid]["toplam_dogru"] += dogru_count
                kazanim_sonuc[kid]["toplam_cevap"] += n

    kazanim_listesi = []
    for kid, ks in kazanim_sonuc.items():
        oran = round(ks["toplam_dogru"] / ks["toplam_cevap"] * 100, 1) if ks["toplam_cevap"] > 0 else 0
        if oran >= 70:
            erisim = "Yeterli"
        elif oran >= 50:
            erisim = "Kismi"
        else:
            erisim = "Yetersiz"

        yorum = ""
        if ks["soru_sayisi"] == 0:
            yorum = "Bu kazanimi olcen soru yok — soru eklenmeli."
            erisim = "Olculemedi"
        elif oran >= 70:
            yorum = f"Ogrencilerin %{oran}'i bu kazanima erisim gostermis. Basarili."
        elif oran >= 50:
            yorum = f"Ogrencilerin %{oran}'i basarili. Ek calisma onerisi yapilin."
        else:
            yorum = f"Basari orani %{oran}. Ogretim yontemi gozden gecirilmeli."

        kazanim_listesi.append({
            "id": kid, "kod": ks["kod"], "aciklama": ks["aciklama"],
            "bloom": ks["bloom"], "soru_sayisi": ks["soru_sayisi"],
            "basari_orani": oran, "erisim": erisim, "yorum": yorum,
        })

    # 17. Uyum matrisi: kazanim x soru
    uyum_matrisi = []
    for kaz in ders_kazanimlar:
        satir = {"kazanim_kod": kaz.kod, "kazanim_id": str(kaz.id), "sorular": []}
        for ss in sinav_sorulari:
            eslesmis = str(kaz.id) in soru_kazanim_map.get(str(ss.soru_id), [])
            satir["sorular"].append({"soru_no": ss.sira, "eslesmis": eslesmis})
        satir["toplam_soru"] = sum(1 for s in satir["sorular"] if s["eslesmis"])
        uyum_matrisi.append(satir)

    # Eslesmemis sorulari bul
    eslesmemis_sorular = []
    for ss in sinav_sorulari:
        if not soru_kazanim_map.get(str(ss.soru_id)):
            eslesmemis_sorular.append(ss.sira)

    # 19. Bloom taksonomisi dagilimi
    bloom_sayim = {}
    for ss in sinav_sorulari:
        bloom = soru_bloom_map.get(str(ss.soru_id), "bilgi")
        bloom_sayim[bloom] = bloom_sayim.get(bloom, 0) + 1

    bloom_toplam = sum(bloom_sayim.values()) or 1
    bloom_dagilimi = []
    bloom_sira = ["bilgi", "kavrama", "uygulama", "analiz", "sentez", "degerlendirme"]
    bloom_labels = {"bilgi": "Bilgi", "kavrama": "Kavrama", "uygulama": "Uygulama", "analiz": "Analiz", "sentez": "Sentez", "degerlendirme": "Degerlendirme"}
    for b in bloom_sira:
        sayi = bloom_sayim.get(b, 0)
        bloom_dagilimi.append({
            "duzey": bloom_labels.get(b, b), "key": b,
            "sayi": sayi, "yuzde": round(sayi / bloom_toplam * 100, 1),
        })

    # Bloom yorumu
    ust_duzey = sum(bloom_sayim.get(b, 0) for b in ["analiz", "sentez", "degerlendirme"])
    alt_duzey = sum(bloom_sayim.get(b, 0) for b in ["bilgi", "kavrama"])
    if ust_duzey / bloom_toplam >= 0.4:
        bloom_yorum = "Ust duzey bilissel beceriler yeterli oranda temsil ediliyor."
    elif ust_duzey / bloom_toplam >= 0.2:
        bloom_yorum = "Ust duzey sorular mevcut ama artirilmasi oneriliyor."
    else:
        bloom_yorum = "Sinav agirlikli olarak alt duzey (bilgi/kavrama) sorulardan olusuyor. Ust duzey sorular eklenmeli."

    return {
        "kazanimlar": kazanim_listesi,
        "uyum_matrisi": uyum_matrisi,
        "eslesmemis_sorular": eslesmemis_sorular,
        "bloom_dagilimi": bloom_dagilimi,
        "bloom_yorum": bloom_yorum,
        "toplam_soru": len(sinav_sorulari),
        "ogrenci_sayisi": n,
    }


@router.get("/analiz/{sinav_id}/kitapcik-karsilastirma")
def kitapcik_karsilastirma(sinav_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    """20. Kitapcik denkleme — A/B/C/D gruplari arasi t-test karsilastirmasi."""
    import math
    sonuclar = db.query(Sonuc).filter_by(sinav_id=str(sinav_id)).all()
    if not sonuclar:
        raise HTTPException(404, "Henuz sonuc yok")

    # Kitapcik bazli gruplama
    gruplar = {}
    for s in sonuclar:
        k = s.kitapcik or "A"
        if k not in gruplar:
            gruplar[k] = []
        gruplar[k].append(s.ham_puan)

    if len(gruplar) < 2:
        return {"gruplar": [], "karsilastirmalar": [], "yorum": "Tek kitapcik kullanilmis, karsilastirma yapilamaz."}

    # Her grup icin istatistikler
    grup_istat = []
    for kit, puanlar in sorted(gruplar.items()):
        n = len(puanlar)
        ort = sum(puanlar) / n
        var = sum((p - ort) ** 2 for p in puanlar) / (n - 1) if n > 1 else 0
        std = math.sqrt(var)
        grup_istat.append({
            "kitapcik": kit, "n": n, "ortalama": round(ort, 2),
            "std_sapma": round(std, 2), "min": round(min(puanlar), 2), "max": round(max(puanlar), 2),
        })

    # Ikili t-test karsilastirmalari
    karsilastirmalar = []
    kitapciklar = sorted(gruplar.keys())
    for i in range(len(kitapciklar)):
        for j in range(i + 1, len(kitapciklar)):
            k1, k2 = kitapciklar[i], kitapciklar[j]
            p1, p2 = gruplar[k1], gruplar[k2]
            n1, n2 = len(p1), len(p2)
            ort1, ort2 = sum(p1) / n1, sum(p2) / n2
            var1 = sum((p - ort1) ** 2 for p in p1) / (n1 - 1) if n1 > 1 else 0
            var2 = sum((p - ort2) ** 2 for p in p2) / (n2 - 1) if n2 > 1 else 0

            # Welch t-test
            se = math.sqrt(var1 / n1 + var2 / n2) if (var1 / n1 + var2 / n2) > 0 else 1
            t_val = (ort1 - ort2) / se
            # Yaklasik p-degeri (normal dagilim yaklasimi)
            try:
                from scipy import stats
                df = ((var1 / n1 + var2 / n2) ** 2) / ((var1 / n1) ** 2 / (n1 - 1) + (var2 / n2) ** 2 / (n2 - 1)) if n1 > 1 and n2 > 1 else n1 + n2 - 2
                p_val = 2 * (1 - stats.t.cdf(abs(t_val), df))
            except Exception:
                p_val = None

            anlamli = p_val is not None and p_val < 0.05
            yorum = f"{k1} vs {k2}: " + (
                f"Anlamli fark var (p={round(p_val, 4)}). Kitapciklar denk degil." if anlamli else
                f"Anlamli fark yok (p={round(p_val, 4) if p_val else '?'}). Kitapciklar denk." if p_val else
                "Hesaplanamadi."
            )

            karsilastirmalar.append({
                "grup1": k1, "grup2": k2,
                "ort1": round(ort1, 2), "ort2": round(ort2, 2),
                "fark": round(ort1 - ort2, 2),
                "t_degeri": round(t_val, 3),
                "p_degeri": round(p_val, 4) if p_val else None,
                "anlamli": anlamli,
                "yorum": yorum,
            })

    genel_yorum = "Tum kitapciklar denk." if all(not k["anlamli"] for k in karsilastirmalar) else "Bazi kitapciklar arasinda anlamli fark var — denkleme gerekli."

    return {"gruplar": grup_istat, "karsilastirmalar": karsilastirmalar, "yorum": genel_yorum}


@router.get("/analiz/ders/{ders_id}/donem-karsilastirma")
def donem_karsilastirma(ders_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    """21. Ayni dersin farkli sinavlarinin karsilastirmasi."""
    import math
    sinavlar = db.query(Sinav).filter_by(ders_id=str(ders_id)).order_by(Sinav.baslangic).all()
    if len(sinavlar) < 1:
        return {"sinavlar": [], "yorum": "Bu derse ait sinav yok."}

    sinav_istatler = []
    for sinav in sinavlar:
        sonuclar = db.query(Sonuc).filter_by(sinav_id=str(sinav.id)).all()
        if not sonuclar:
            continue
        puanlar = [s.ham_puan for s in sonuclar]
        n = len(puanlar)
        ort = sum(puanlar) / n
        var = sum((p - ort) ** 2 for p in puanlar) / n if n > 0 else 0
        std = math.sqrt(var)
        sinav_istatler.append({
            "sinav_id": str(sinav.id), "ad": sinav.ad, "sinav_turu": sinav.sinav_turu,
            "tarih": sinav.baslangic.strftime('%d.%m.%Y') if sinav.baslangic else None,
            "n": n, "ortalama": round(ort, 2), "std_sapma": round(std, 2),
            "min": round(min(puanlar), 2), "max": round(max(puanlar), 2),
            "basari_orani": round(sum(1 for p in puanlar if p >= sinav.tam_puan * 0.5) / n * 100, 1),
        })

    if len(sinav_istatler) < 2:
        yorum = "Karsilastirma icin en az 2 sinav gerekli."
    else:
        ortalar = [s["ortalama"] for s in sinav_istatler]
        trend = "yukseliyor" if ortalar[-1] > ortalar[0] else "dusuyor" if ortalar[-1] < ortalar[0] else "sabit"
        yorum = f"Sinav ortalamasi {trend}. Ilk sinav: {ortalar[0]}, son sinav: {ortalar[-1]}."

    return {"sinavlar": sinav_istatler, "yorum": yorum}


@router.get("/analiz/ders-haritasi")
def ders_basari_haritasi(db: Session = Depends(get_db), _=Depends(get_user)):
    """22. Tum derslerin genel basari ozeti."""
    import math
    dersler = db.query(Ders).filter_by(aktif=True).all()
    harita = []
    for ders in dersler:
        sinavlar = db.query(Sinav).filter_by(ders_id=str(ders.id)).all()
        if not sinavlar:
            continue
        tum_puanlar = []
        sinav_sayisi = 0
        for sinav in sinavlar:
            sonuclar = db.query(Sonuc).filter_by(sinav_id=str(sinav.id)).all()
            if sonuclar:
                sinav_sayisi += 1
                tum_puanlar.extend([s.ham_puan for s in sonuclar])

        if not tum_puanlar:
            continue

        n = len(tum_puanlar)
        ort = sum(tum_puanlar) / n
        harita.append({
            "ders_id": str(ders.id), "ders_ad": ders.ad,
            "sinav_sayisi": sinav_sayisi, "ogrenci_sayisi": n,
            "ortalama": round(ort, 2),
            "basari_orani": round(sum(1 for p in tum_puanlar if p >= 50) / n * 100, 1),
        })

    harita.sort(key=lambda x: x["basari_orani"], reverse=True)
    return {"dersler": harita}


# ══════════════════════════════════════════════════════════════════
#  SORU SABLON + TOPLU YUKLEME
# ══════════════════════════════════════════════════════════════════

@router.get("/soru-sablon")
def soru_sablon_indir(_=Depends(get_user)):
    """Soru yukleme Excel sablonu."""
    import openpyxl
    from urllib.parse import quote
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sorular"
    basliklar = ["Soru Metni", "A", "B", "C", "D", "E", "Dogru Cevap", "Soru Tipi", "Bilgisel Duzey", "Zorluk", "Cevaplama Suresi", "Anahtar Kelimeler"]
    for i, h in enumerate(basliklar, 1):
        ws.cell(row=1, column=i, value=h).font = openpyxl.styles.Font(bold=True)
    ws.append(["Hucre bolunmesinde hangi organel gorev alir?", "Mitokondri", "Ribozom", "Sentrozom", "Golgi", "Lizozom", "C", "tek_dogru", "Bilgi", "orta", 60, "hucre,bolunme"])
    for col in ws.columns:
        ws.column_dimensions[col[0].column_letter].width = 18

    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                             headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote('soru_sablonu.xlsx')}"})


@router.post("/sorular/toplu-yukle/{konu_id}")
async def sorular_toplu_yukle(konu_id: UUID, dosya: UploadFile = File(...), db: Session = Depends(get_db), me=Depends(get_yetkili_user)):
    """Excel'den toplu soru yukle."""
    import openpyxl
    icerik = await dosya.read()
    wb = openpyxl.load_workbook(io.BytesIO(icerik), read_only=True, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if len(rows) < 2:
        raise HTTPException(400, "Excel'de veri yok")

    harfler = "ABCDE"
    eklenen = 0
    for row in rows[1:]:
        if not row or not row[0]: continue
        soru_metni = str(row[0]).strip()
        secenekler = []
        dogru_cevap = str(row[6] or "A").strip().upper() if len(row) > 6 else "A"

        for i in range(5):
            if len(row) > i + 1 and row[i + 1]:
                secenekler.append({
                    "metin": str(row[i + 1]).strip(),
                    "dogru": harfler[i] == dogru_cevap,
                    "sira": i,
                })

        s = Soru(
            konu_id=str(konu_id), soru_metni=soru_metni,
            soru_tipi=str(row[7] or "tek_dogru").strip() if len(row) > 7 else "tek_dogru",
            bilgisel_duzey=str(row[8] or "Bilgi").strip() if len(row) > 8 else "Bilgi",
            zorluk=str(row[9] or "orta").strip() if len(row) > 9 else "orta",
            cevaplama_suresi=int(float(str(row[10] or 60))) if len(row) > 10 else 60,
            anahtar_kelimeler=str(row[11] or "").split(",") if len(row) > 11 and row[11] else [],
            olusturan_id=str(me.id),
        )
        db.add(s); db.flush()
        for sc in secenekler:
            db.add(SoruSecenegi(soru_id=str(s.id), secenek_metni=sc["metin"], dogru=sc["dogru"], sira=sc["sira"]))
        eklenen += 1

    db.commit()
    return {"eklenen": eklenen}


# ══════════════════════════════════════════════════════════════════
#  DERS OGRENIM KAZANIMLARI
# ══════════════════════════════════════════════════════════════════

@router.get("/dersler/{ders_id}/kazanimlar")
def kazanim_listele(ders_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    kazanimlar = db.query(DersKazanim).filter_by(ders_id=str(ders_id)).order_by(DersKazanim.sira).all()
    sonuc = []
    for k in kazanimlar:
        konu_ids = [str(kk.konu_id) for kk in db.query(KazanimKonu).filter_by(kazanim_id=str(k.id)).all()]
        sonuc.append({
            "id": str(k.id), "kod": k.kod, "aciklama": k.aciklama,
            "bloom_duzeyi": k.bloom_duzeyi, "sira": k.sira,
            "konu_ids": konu_ids,
        })
    return sonuc

@router.post("/dersler/{ders_id}/kazanimlar")
def kazanim_ekle(ders_id: UUID, veri: dict, db: Session = Depends(get_db), _=Depends(get_yetkili_user)):
    mevcut = db.query(DersKazanim).filter_by(ders_id=str(ders_id)).count()
    kod = veri.get("kod") or f"OK-{mevcut + 1}"
    k = DersKazanim(
        ders_id=str(ders_id), kod=kod, aciklama=veri["aciklama"],
        bloom_duzeyi=veri.get("bloom_duzeyi", "Bilgi"), sira=veri.get("sira", mevcut + 1),
    )
    db.add(k); db.flush()
    for konu_id in veri.get("konu_ids", []):
        db.add(KazanimKonu(kazanim_id=str(k.id), konu_id=konu_id))
    db.commit()
    return {"id": str(k.id), "kod": k.kod}

@router.put("/kazanimlar/{kazanim_id}")
def kazanim_guncelle(kazanim_id: UUID, veri: dict, db: Session = Depends(get_db), _=Depends(get_yetkili_user)):
    k = db.query(DersKazanim).filter_by(id=str(kazanim_id)).first()
    if not k: raise HTTPException(404, "Kazanim bulunamadi")
    if "kod" in veri: k.kod = veri["kod"]
    if "aciklama" in veri: k.aciklama = veri["aciklama"]
    if "bloom_duzeyi" in veri: k.bloom_duzeyi = veri["bloom_duzeyi"]
    if "sira" in veri: k.sira = veri["sira"]
    if "konu_ids" in veri:
        db.query(KazanimKonu).filter_by(kazanim_id=str(kazanim_id)).delete()
        for konu_id in veri["konu_ids"]:
            db.add(KazanimKonu(kazanim_id=str(kazanim_id), konu_id=konu_id))
    db.commit()
    return {"ok": True}

@router.post("/dersler/{ders_id}/kazanimlar/toplu")
def kazanim_toplu_ekle(ders_id: UUID, veri: dict, db: Session = Depends(get_db), _=Depends(get_yetkili_user)):
    """Yapistirilan metinden kazanimlari otomatik olustur.
    veri: { metin: str } veya { kazanimlar: [{aciklama, bloom_duzeyi}] }
    Metin formati:
      - 'OK-1: Aciklama [Bilgi]'
      - '1. Aciklama (Kavrama)'
      - Her satir bir kazanim
    """
    import re
    kazanimlar_list = veri.get("kazanimlar")
    if not kazanimlar_list:
        metin = veri.get("metin", "").strip()
        if not metin:
            raise HTTPException(400, "Metin veya kazanimlar listesi gerekli")
        satirlar = [s.strip() for s in metin.split("\n") if s.strip()]
        kazanimlar_list = []
        bloom_keywords = {
            "bilgi": "Bilgi", "kavrama": "Kavrama", "uygulama": "Uygulama",
            "analiz": "Analiz", "sentez": "Sentez", "degerlendirme": "Degerlendirme",
            "knowledge": "Bilgi", "comprehension": "Kavrama", "application": "Uygulama",
            "analysis": "Analiz", "synthesis": "Sentez", "evaluation": "Degerlendirme",
        }
        for satir in satirlar:
            bloom = "Bilgi"
            # [Bilgi] veya (Kavrama) formatini bul
            bm = re.search(r"[\[\(]([\wğüşıöçĞÜŞİÖÇ]+)[\]\)]$", satir)
            if bm:
                b_text = bm.group(1).lower()
                for key, val in bloom_keywords.items():
                    if key in b_text:
                        bloom = val
                        break
                satir = satir[:bm.start()].strip()
            # OK-1: veya 1. veya 1) prefix
            m = re.match(r"^(?:OK|ÖK|LO)\s*[-.]?\s*(\d+)\s*[:.\-–]\s*(.+)", satir, re.IGNORECASE)
            if m:
                kazanimlar_list.append({"aciklama": m.group(2).strip(), "bloom_duzeyi": bloom})
                continue
            m2 = re.match(r"^(\d+)\s*[.):\-–]\s*(.+)", satir)
            if m2:
                kazanimlar_list.append({"aciklama": m2.group(2).strip(), "bloom_duzeyi": bloom})
                continue
            if satir:
                kazanimlar_list.append({"aciklama": satir, "bloom_duzeyi": bloom})

    mevcut = db.query(DersKazanim).filter_by(ders_id=str(ders_id)).count()
    eklenen = 0
    for i, k in enumerate(kazanimlar_list):
        sira = mevcut + i + 1
        db.add(DersKazanim(
            ders_id=str(ders_id), kod=k.get("kod") or f"OK-{sira}",
            aciklama=k["aciklama"], bloom_duzeyi=k.get("bloom_duzeyi", "Bilgi"), sira=sira,
        ))
        eklenen += 1
    db.commit()
    return {"eklenen": eklenen}

@router.delete("/kazanimlar/{kazanim_id}")
def kazanim_sil(kazanim_id: UUID, db: Session = Depends(get_db), _=Depends(get_yetkili_user)):
    db.query(KazanimKonu).filter_by(kazanim_id=str(kazanim_id)).delete()
    db.query(DersKazanim).filter_by(id=str(kazanim_id)).delete()
    db.commit()
    return {"ok": True}


# ══════════════════════════════════════════════════════════════════
#  OTOMATIK SINAV KAGIDI OLUSTURMA
# ══════════════════════════════════════════════════════════════════

@router.post("/sinavlar/{sinav_id}/otomatik-olustur")
def otomatik_sinav_olustur(sinav_id: UUID, veri: dict, db: Session = Depends(get_db), _=Depends(get_yetkili_user)):
    """
    Otomatik soru secimi ile sinav kagidi olusturur.
    veri: {
        grup_sayisi: int (kac farkli grup/kitapcik, default 1),
        dagilim: "esit" | "oranli" | "manuel",
        zorluk_dagilimi: { kolay: 30, orta: 50, zor: 20 } (yuzde),
        konu_agirliklari: { konu_id: soru_sayisi } (sadece manuel icin),
    }
    """
    import random

    sinav = db.query(Sinav).filter_by(id=str(sinav_id)).first()
    if not sinav:
        raise HTTPException(404, "Sinav bulunamadi")

    soru_sayisi = sinav.soru_sayisi or 0
    if soru_sayisi == 0:
        raise HTTPException(400, "Sinav soru sayisi belirlenmemis")

    grup_sayisi = veri.get("grup_sayisi", 1)
    dagilim = veri.get("dagilim", "esit")
    zorluk_dag = veri.get("zorluk_dagilimi", {"kolay": 25, "orta": 50, "zor": 25})

    # Dersin tum konularini bul
    ders = db.query(Ders).filter_by(id=str(sinav.ders_id)).first()
    if not ders:
        raise HTTPException(404, "Ders bulunamadi")

    konular = db.query(Konu).filter_by(ders_id=str(ders.id)).order_by(Konu.sira).all()
    if not konular:
        raise HTTPException(400, "Dersin konusu yok, once konu ekleyin")

    # Tum sorulari topla
    tum_sorular = []
    for konu in konular:
        sorular = db.query(Soru).filter_by(konu_id=str(konu.id)).all()
        for s in sorular:
            tum_sorular.append({"id": str(s.id), "konu_id": str(konu.id), "zorluk": s.zorluk or "orta", "konu_ad": konu.ad})

    if len(tum_sorular) < soru_sayisi:
        raise HTTPException(400, f"Yeterli soru yok: {len(tum_sorular)} mevcut, {soru_sayisi} gerekli")

    # Onceki atanmis sorulari temizle
    db.query(SinavSorusu).filter_by(sinav_id=str(sinav_id)).delete()
    db.query(SinavPlani).filter_by(sinav_id=str(sinav_id)).delete()

    gruplar = []
    kitapcik_harfler = "ABCDEFGH"

    for g in range(grup_sayisi):
        kitapcik = kitapcik_harfler[g] if g < len(kitapcik_harfler) else str(g + 1)

        if dagilim == "manuel" and veri.get("konu_agirliklari"):
            # Manuel: konu bazli soru sayilari belirli
            secilen = []
            for konu_id, adet in veri["konu_agirliklari"].items():
                konu_sorulari = [s for s in tum_sorular if s["konu_id"] == konu_id]
                random.shuffle(konu_sorulari)
                secilen.extend(konu_sorulari[:adet])
        else:
            # Zorluk dagilimina gore sec
            kolay_n = int(soru_sayisi * zorluk_dag.get("kolay", 25) / 100)
            zor_n = int(soru_sayisi * zorluk_dag.get("zor", 25) / 100)
            orta_n = soru_sayisi - kolay_n - zor_n

            kolay_havuz = [s for s in tum_sorular if s["zorluk"] in ("kolay", "cok_kolay")]
            orta_havuz = [s for s in tum_sorular if s["zorluk"] == "orta"]
            zor_havuz = [s for s in tum_sorular if s["zorluk"] in ("zor", "cok_zor")]

            random.shuffle(kolay_havuz)
            random.shuffle(orta_havuz)
            random.shuffle(zor_havuz)

            secilen = kolay_havuz[:kolay_n] + orta_havuz[:orta_n] + zor_havuz[:zor_n]

            # Eksik kalirsa rastgele tamamla
            if len(secilen) < soru_sayisi:
                kalanlar = [s for s in tum_sorular if s not in secilen]
                random.shuffle(kalanlar)
                secilen.extend(kalanlar[:soru_sayisi - len(secilen)])

        # Grup icinde siralama karistir
        random.shuffle(secilen)
        secilen = secilen[:soru_sayisi]

        # DB'ye kaydet
        for sira, s in enumerate(secilen, 1):
            db.add(SinavSorusu(sinav_id=str(sinav_id), soru_id=s["id"], sira=sira, kitapcik=kitapcik))

        # Konu bazli plan olustur
        konu_sayac = {}
        for s in secilen:
            konu_sayac[s["konu_id"]] = konu_sayac.get(s["konu_id"], 0) + 1

        gruplar.append({
            "kitapcik": kitapcik,
            "soru_sayisi": len(secilen),
            "konu_dagilimi": {s["konu_ad"]: konu_sayac.get(s["konu_id"], 0) for s in secilen},
            "zorluk_dagilimi": {
                "kolay": sum(1 for s in secilen if s["zorluk"] in ("kolay", "cok_kolay")),
                "orta": sum(1 for s in secilen if s["zorluk"] == "orta"),
                "zor": sum(1 for s in secilen if s["zorluk"] in ("zor", "cok_zor")),
            }
        })

    # Sinav planini kaydet (ilk grubun konu dagilimi)
    if gruplar:
        ilk = gruplar[0]
        for konu in konular:
            adet = 0
            for s_item in db.query(SinavSorusu).filter_by(sinav_id=str(sinav_id), kitapcik=gruplar[0]["kitapcik"]).all():
                soru = db.query(Soru).filter_by(id=str(s_item.soru_id)).first()
                if soru and str(soru.konu_id) == str(konu.id):
                    adet += 1
            if adet > 0:
                db.add(SinavPlani(sinav_id=str(sinav_id), konu_id=str(konu.id), hafta=konu.hafta, saat=konu.saat, gerekli_soru=adet, secilen_soru=adet))

    sinav.durum = "yaratildi"
    sinav.kitapcik_turu = kitapcik_harfler[:grup_sayisi] if grup_sayisi <= len(kitapcik_harfler) else f"{grup_sayisi} grup"
    db.commit()

    return {
        "sinav_id": str(sinav_id),
        "grup_sayisi": grup_sayisi,
        "toplam_atanan_soru": db.query(SinavSorusu).filter_by(sinav_id=str(sinav_id)).count(),
        "gruplar": gruplar,
    }


@router.get("/sinavlar/{sinav_id}/kagit")
def sinav_kagidi(sinav_id: UUID, kitapcik: str = "A", db: Session = Depends(get_db), _=Depends(get_user)):
    """Sinav kagidi - soru listesi sirasiyla."""
    sinav = db.query(Sinav).filter_by(id=str(sinav_id)).first()
    if not sinav:
        raise HTTPException(404, "Sinav bulunamadi")

    sinav_sorulari = db.query(SinavSorusu).filter_by(sinav_id=str(sinav_id), kitapcik=kitapcik).order_by(SinavSorusu.sira).all()
    if not sinav_sorulari:
        sinav_sorulari = db.query(SinavSorusu).filter_by(sinav_id=str(sinav_id)).order_by(SinavSorusu.sira).all()

    sorular = []
    for ss in sinav_sorulari:
        soru = db.query(Soru).filter_by(id=str(ss.soru_id)).first()
        if not soru:
            continue
        secenekler = db.query(SoruSecenegi).filter_by(soru_id=str(soru.id)).order_by(SoruSecenegi.sira).all()
        konu = db.query(Konu).filter_by(id=str(soru.konu_id)).first()
        sorular.append({
            "sira": ss.sira,
            "soru_metni": soru.soru_metni,
            "soru_tipi": soru.soru_tipi,
            "zorluk": soru.zorluk,
            "konu": konu.ad if konu else None,
            "secenekler": [{"harf": "ABCDE"[i], "metin": sc.secenek_metni, "dogru": sc.dogru} for i, sc in enumerate(secenekler)],
        })

    # Kitapciklari listele
    kitapciklar = list(set(ss.kitapcik for ss in db.query(SinavSorusu).filter_by(sinav_id=str(sinav_id)).all() if ss.kitapcik))
    kitapciklar.sort()

    return {
        "sinav_ad": sinav.ad,
        "kitapcik": kitapcik,
        "kitapciklar": kitapciklar,
        "soru_sayisi": len(sorular),
        "sorular": sorular,
    }


@router.get("/sinavlar/{sinav_id}/plan-detay")
def sinav_plan_detay(sinav_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    """Sinav plani detayi - konu bazli dagilim."""
    planlar = db.query(SinavPlani).filter_by(sinav_id=str(sinav_id)).all()
    sonuc = []
    for p in planlar:
        konu = db.query(Konu).filter_by(id=str(p.konu_id)).first()
        sonuc.append({
            "konu_id": str(p.konu_id), "konu_ad": konu.ad if konu else "?",
            "hafta": p.hafta, "saat": p.saat,
            "gerekli_soru": p.gerekli_soru, "secilen_soru": p.secilen_soru,
        })
    return sonuc


# ══════════════════════════════════════════════════════════════════
#  DERS BAZLI SINAV LISTELEME
# ══════════════════════════════════════════════════════════════════

@router.get("/dersler/{ders_id}/sinavlar")
def ders_sinavlari(ders_id: UUID, db: Session = Depends(get_db), _=Depends(get_user)):
    sinavlar = db.query(Sinav).filter_by(ders_id=str(ders_id)).order_by(Sinav.baslangic.desc()).all()
    sonuc = []
    for s in sinavlar:
        sonuc_sayisi = db.query(Sonuc).filter_by(sinav_id=str(s.id)).count()
        sonuc.append({
            "id": str(s.id), "ad": s.ad, "sinav_turu": s.sinav_turu,
            "sinav_kategorisi": s.sinav_kategorisi, "durum": s.durum,
            "soru_sayisi": s.soru_sayisi, "tam_puan": s.tam_puan,
            "baslangic": s.baslangic.isoformat() if s.baslangic else None,
            "bitis": s.bitis.isoformat() if s.bitis else None,
            "sonuc_var": sonuc_sayisi > 0,
            "ogrenci_sayisi": sonuc_sayisi,
        })
    return sonuc


# ══════════════════════════════════════════════════════════════════
#  PDF SINAV KAGIDI
# ══════════════════════════════════════════════════════════════════

@router.get("/sinavlar/{sinav_id}/pdf")
def sinav_pdf(sinav_id: UUID, kitapcik: str = "A", cevap_anahtari: bool = False,
              db: Session = Depends(get_db), _=Depends(get_user)):
    """Sinav kagidini PDF olarak indir."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib import colors
    import re as re_mod

    sinav = db.query(Sinav).filter_by(id=str(sinav_id)).first()
    if not sinav:
        raise HTTPException(404, "Sinav bulunamadi")

    # Soruları çek
    sinav_sorulari = db.query(SinavSorusu).filter_by(
        sinav_id=str(sinav_id), kitapcik=kitapcik
    ).order_by(SinavSorusu.sira).all()
    if not sinav_sorulari:
        sinav_sorulari = db.query(SinavSorusu).filter_by(
            sinav_id=str(sinav_id)
        ).order_by(SinavSorusu.sira).all()

    if not sinav_sorulari:
        raise HTTPException(404, "Bu sinav icin soru bulunamadi")

    sorular = []
    for ss in sinav_sorulari:
        soru = db.query(Soru).filter_by(id=str(ss.soru_id)).first()
        if not soru:
            continue
        secenekler = db.query(SoruSecenegi).filter_by(
            soru_id=str(soru.id)
        ).order_by(SoruSecenegi.sira).all()
        # Kazanim kodlarini al
        kazanim_kodlari = []
        for sk in db.query(SoruKazanim).filter_by(soru_id=str(soru.id)).all():
            kaz = db.query(DersKazanim).filter_by(id=str(sk.kazanim_id)).first()
            if kaz:
                kazanim_kodlari.append(kaz.kod)

        sorular.append({
            "sira": ss.sira,
            "soru_metni": soru.soru_metni,
            "kazanim_kodlari": kazanim_kodlari,
            "secenekler": [
                {"harf": "ABCDE"[i], "metin": sc.secenek_metni, "dogru": sc.dogru}
                for i, sc in enumerate(secenekler)
            ],
        })

    # Ders bilgisi
    ders = db.query(Ders).filter_by(id=str(sinav.ders_id)).first()
    ders_ad = ders.ad if ders else ""

    # PDF oluştur
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        topMargin=20*mm, bottomMargin=15*mm,
        leftMargin=18*mm, rightMargin=18*mm,
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='SinavBaslik', fontSize=16, fontName='Helvetica-Bold',
        alignment=TA_CENTER, spaceAfter=4*mm,
    ))
    styles.add(ParagraphStyle(
        name='SinavAltBaslik', fontSize=11, fontName='Helvetica',
        alignment=TA_CENTER, spaceAfter=6*mm, textColor=colors.HexColor('#555555'),
    ))
    styles.add(ParagraphStyle(
        name='SoruMetni', fontSize=11, fontName='Helvetica',
        leading=15, spaceAfter=2*mm, leftIndent=6*mm,
    ))
    styles.add(ParagraphStyle(
        name='SoruNo', fontSize=11, fontName='Helvetica-Bold',
        spaceAfter=1*mm,
    ))
    styles.add(ParagraphStyle(
        name='Secenek', fontSize=10, fontName='Helvetica',
        leading=14, leftIndent=12*mm,
    ))
    styles.add(ParagraphStyle(
        name='SecenekDogru', fontSize=10, fontName='Helvetica-Bold',
        leading=14, leftIndent=12*mm, textColor=colors.HexColor('#10b981'),
    ))

    elements = []

    # Başlık
    elements.append(Paragraph(sinav.ad, styles['SinavBaslik']))
    alt_bilgi = f"{ders_ad}"
    if kitapcik:
        alt_bilgi += f" | Kitapcik {kitapcik}"
    alt_bilgi += f" | {len(sorular)} Soru | {sinav.tam_puan} Puan"
    elements.append(Paragraph(alt_bilgi, styles['SinavAltBaslik']))

    # Öğrenci bilgi alanı
    bilgi_data = [
        ['Ad Soyad:', '____________________________', 'Ogrenci No:', '________________'],
        ['Tarih:', sinav.baslangic.strftime('%d.%m.%Y') if sinav.baslangic else '___/___/______', 'Imza:', '________________'],
    ]
    bilgi_tablo = Table(bilgi_data, colWidths=[22*mm, 60*mm, 22*mm, 50*mm])
    bilgi_tablo.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4*mm),
        ('TOPPADDING', (0, 0), (-1, -1), 2*mm),
    ]))
    elements.append(bilgi_tablo)
    elements.append(Spacer(1, 4*mm))

    # Çizgi
    line_data = [['_' * 100]]
    line_table = Table(line_data, colWidths=[174*mm])
    line_table.setStyle(TableStyle([
        ('FONTSIZE', (0, 0), (-1, -1), 6),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#cccccc')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4*mm),
    ]))
    elements.append(line_table)

    # Sorular
    for s in sorular:
        metin = re_mod.sub(r'<[^>]+>', '', s["soru_metni"])

        # Kazanim kodu varsa soru numarasinin yanina ekle
        kazanim_str = ""
        if s.get("kazanim_kodlari"):
            kazanim_str = f"  <font color='#6b7280' size='9'>(O.K. {', '.join(s['kazanim_kodlari'])})</font>"
        elements.append(Paragraph(f"Soru {s['sira']}.{kazanim_str}", styles['SoruNo']))
        elements.append(Paragraph(metin, styles['SoruMetni']))

        for sec in s["secenekler"]:
            if cevap_anahtari and sec["dogru"]:
                elements.append(Paragraph(
                    f"<b>{sec['harf']})</b> {sec['metin']}",
                    styles['SecenekDogru']
                ))
            else:
                elements.append(Paragraph(
                    f"<b>{sec['harf']})</b> {sec['metin']}",
                    styles['Secenek']
                ))

        elements.append(Spacer(1, 4*mm))

    # Cevap anahtarı sayfası
    if cevap_anahtari:
        elements.append(PageBreak())
        elements.append(Paragraph("CEVAP ANAHTARI", styles['SinavBaslik']))
        elements.append(Paragraph(f"{sinav.ad} - Kitapcik {kitapcik}", styles['SinavAltBaslik']))

        cevap_data = [['Soru', 'Dogru Cevap']]
        for s in sorular:
            dogru = next((sec["harf"] for sec in s["secenekler"] if sec["dogru"]), "-")
            cevap_data.append([str(s["sira"]), dogru])

        ct = Table(cevap_data, colWidths=[30*mm, 40*mm])
        ct.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#6366f1')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e5e7eb')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9fafb')]),
            ('TOPPADDING', (0, 0), (-1, -1), 3*mm),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3*mm),
        ]))
        elements.append(ct)

    doc.build(elements)
    buffer.seek(0)

    from urllib.parse import quote
    filename = f"{sinav.ad.replace(' ', '_')}_Kitapcik_{kitapcik}.pdf"
    return StreamingResponse(
        buffer, media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"}
    )


# ══════════════════════════════════════════════════════════════════
#  BILDIRIMLER
# ══════════════════════════════════════════════════════════════════

@router.get("/bildirimler")
def bildirimler_listele(db: Session = Depends(get_db), me=Depends(get_user)):
    """Kullanicinin bildirimlerini getir (son 50)."""
    items = db.query(Bildirim).filter_by(
        kullanici_id=str(me.id)
    ).order_by(Bildirim.olusturma.desc()).limit(50).all()
    return [
        {
            "id": str(b.id), "baslik": b.baslik, "mesaj": b.mesaj,
            "tip": b.tip, "okundu": b.okundu, "link": b.link,
            "olusturma": b.olusturma.isoformat() if b.olusturma else None,
        }
        for b in items
    ]


@router.get("/bildirimler/okunmamis-sayisi")
def okunmamis_bildirim_sayisi(db: Session = Depends(get_db), me=Depends(get_user)):
    sayi = db.query(Bildirim).filter_by(
        kullanici_id=str(me.id), okundu=False
    ).count()
    return {"sayi": sayi}


@router.put("/bildirimler/{bildirim_id}/oku")
def bildirim_oku(bildirim_id: UUID, db: Session = Depends(get_db), me=Depends(get_user)):
    b = db.query(Bildirim).filter_by(id=str(bildirim_id), kullanici_id=str(me.id)).first()
    if not b:
        raise HTTPException(404, "Bildirim bulunamadi")
    b.okundu = True
    db.commit()
    return {"ok": True}


@router.put("/bildirimler/tumunu-oku")
def tumunu_oku(db: Session = Depends(get_db), me=Depends(get_user)):
    db.query(Bildirim).filter_by(
        kullanici_id=str(me.id), okundu=False
    ).update({"okundu": True})
    db.commit()
    return {"ok": True}


# ══════════════════════════════════════════════════════════════════
#  DOCX SINAV KAGIDI (SABLON BAZLI)
# ══════════════════════════════════════════════════════════════════

@router.get("/sinavlar/{sinav_id}/docx")
def sinav_docx(sinav_id: UUID, kitapcik: str = "A", db: Session = Depends(get_db), _=Depends(get_user)):
    """Sinav kagidini Medipol sablon uzerinde DOCX olarak olustur."""
    import os, copy, re as re_mod
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from lxml import etree

    sinav = db.query(Sinav).filter_by(id=str(sinav_id)).first()
    if not sinav:
        raise HTTPException(404, "Sinav bulunamadi")

    # Sorulari cek
    sinav_sorulari = db.query(SinavSorusu).filter_by(
        sinav_id=str(sinav_id), kitapcik=kitapcik
    ).order_by(SinavSorusu.sira).all()
    if not sinav_sorulari:
        sinav_sorulari = db.query(SinavSorusu).filter_by(
            sinav_id=str(sinav_id)
        ).order_by(SinavSorusu.sira).all()
    if not sinav_sorulari:
        raise HTTPException(404, "Bu sinav icin soru bulunamadi")

    sorular = []
    for ss in sinav_sorulari:
        soru = db.query(Soru).filter_by(id=str(ss.soru_id)).first()
        if not soru:
            continue
        secenekler = db.query(SoruSecenegi).filter_by(
            soru_id=str(soru.id)
        ).order_by(SoruSecenegi.sira).all()
        # Kazanim kodlari
        kazanim_kodlari = []
        for sk in db.query(SoruKazanim).filter_by(soru_id=str(soru.id)).all():
            kaz = db.query(DersKazanim).filter_by(id=str(sk.kazanim_id)).first()
            if kaz:
                kazanim_kodlari.append(kaz.kod)
        sorular.append({
            "sira": ss.sira,
            "soru_metni": soru.soru_metni,
            "kazanim_kodlari": kazanim_kodlari,
            "secenekler": [
                {"harf": "ABCDE"[i], "metin": sc.secenek_metni}
                for i, sc in enumerate(secenekler)
            ],
        })

    # Ders ve diger bilgiler
    ders = db.query(Ders).filter_by(id=str(sinav.ders_id)).first()
    ders_ad = ders.ad if ders else ""

    # Sablonu yukle
    sablon_yolu = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "sinav_sablonu.docx")
    if not os.path.exists(sablon_yolu):
        raise HTTPException(500, "Sinav sablonu bulunamadi")

    doc = DocxDocument(sablon_yolu)

    # Header'daki sinav turunu doldur
    for section in doc.sections:
        for p in section.header.paragraphs:
            if '………' in p.text:
                for run in p.runs:
                    if '………' in run.text:
                        sinav_turu_label = {
                            'donem_ici': 'ARA SINAV', 'final': 'FİNAL',
                            'butunleme': 'BÜTÜNLEME', 'hazir_bulunurluk': 'HAZIR BULUNURLUK',
                            'gelisim': 'GELİŞİM',
                        }.get(sinav.sinav_turu, 'SINAV')
                        run.text = run.text.replace('………………………………', sinav_turu_label)

    # Textbox'lardaki bilgileri doldur
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
          'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'}
    for txbx in doc.element.body.findall('.//wps:txbx', ns):
        for p in txbx.findall('.//w:p', ns):
            full_text = ''.join(p.itertext())
            if 'Dersin Adı' in full_text or 'Course Name' in full_text:
                # Son run'a ders adini ekle
                runs = p.findall('.//w:r', ns)
                if runs:
                    last_t = runs[-1].find('w:t', ns)
                    if last_t is not None:
                        last_t.text = (last_t.text or '') + ' ' + ders_ad
            elif 'Sınav Tarihi' in full_text or 'Date' in full_text:
                tarih = sinav.baslangic.strftime('%d.%m.%Y %H:%M') if sinav.baslangic else ''
                runs = p.findall('.//w:r', ns)
                if runs:
                    last_t = runs[-1].find('w:t', ns)
                    if last_t is not None:
                        last_t.text = (last_t.text or '') + ' ' + tarih

    # Mevcut bos paragraf(:) lerini sil ve sorulari ekle
    body = doc.element.body
    # Son paragraflardan ':' olanlari bul
    bos_paragraflar = []
    for p_elem in body.findall('w:p', ns):
        text = ''.join(p_elem.itertext()).strip()
        if text == ':' or text == '':
            bos_paragraflar.append(p_elem)

    # Bos paragraflari sil
    for bp in bos_paragraflar:
        body.remove(bp)

    # Sorulari ekle (body'nin sonuna, sectPr'den once)
    sect_pr = body.find('w:sectPr', ns)

    for s in sorular:
        metin = re_mod.sub(r'<[^>]+>', '', s["soru_metni"])
        kazanim_str = f"  (Ö.K. {', '.join(s['kazanim_kodlari'])})" if s.get("kazanim_kodlari") else ""

        # Soru numarasi + metin
        soru_p = doc.add_paragraph()
        soru_p.paragraph_format.space_before = Pt(8)
        soru_p.paragraph_format.space_after = Pt(4)
        run_no = soru_p.add_run(f"{s['sira']}. ")
        run_no.bold = True
        run_no.font.size = Pt(11)
        run_metin = soru_p.add_run(metin)
        run_metin.font.size = Pt(11)
        if kazanim_str:
            run_kaz = soru_p.add_run(kazanim_str)
            run_kaz.font.size = Pt(8)
            run_kaz.font.color.rgb = None  # gri

        # Secenekler
        for sec in s["secenekler"]:
            sec_p = doc.add_paragraph()
            sec_p.paragraph_format.space_before = Pt(1)
            sec_p.paragraph_format.space_after = Pt(1)
            sec_p.paragraph_format.left_indent = Cm(1)
            run_harf = sec_p.add_run(f"{sec['harf']}) ")
            run_harf.bold = True
            run_harf.font.size = Pt(10)
            run_sec = sec_p.add_run(sec["metin"])
            run_sec.font.size = Pt(10)

    # sectPr'yi sona tasi
    if sect_pr is not None:
        body.remove(sect_pr)
        body.append(sect_pr)

    # Kaydet
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    from urllib.parse import quote
    filename = f"{sinav.ad.replace(' ', '_')}_Kitapcik_{kitapcik}.docx"
    return StreamingResponse(
        buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"}
    )
